"""Build a simple Word manuscript from docs/manuscript_draft_v0.md.

The converter is intentionally narrow: it supports the Markdown patterns used in
the manuscript draft (headings, paragraphs, bullet/numbered lists, bold spans,
inline code and fenced code blocks) and applies a clean journal-manuscript style.
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "docs" / "manuscript_draft_v0.md"
OUT = ROOT / "docs" / "manuscript_draft_v0.docx"


def set_spacing(style, *, before=0, after=6, line=1.15):
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_font(style, *, name="Calibri", size=11, color="000000", bold=None):
    font = style.font
    font.name = name
    font.size = Pt(size)
    font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def add_bottom_border(paragraph, color="DADCE0", size="6"):
    p = paragraph._p
    ppr = p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def add_shading(paragraph, fill="F4F6F9"):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def shade_cell(cell, fill="E8EEF5"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_font(styles["Normal"], name="Calibri", size=11, color="000000")
    set_spacing(styles["Normal"], before=0, after=6, line=1.15)

    set_font(styles["Title"], name="Calibri", size=20, color="0B2545", bold=True)
    set_spacing(styles["Title"], before=0, after=12, line=1.15)

    set_font(styles["Heading 1"], name="Calibri", size=16, color="2E74B5", bold=True)
    set_spacing(styles["Heading 1"], before=14, after=6, line=1.15)
    set_font(styles["Heading 2"], name="Calibri", size=13, color="2E74B5", bold=True)
    set_spacing(styles["Heading 2"], before=12, after=5, line=1.15)
    set_font(styles["Heading 3"], name="Calibri", size=12, color="1F4D78", bold=True)
    set_spacing(styles["Heading 3"], before=10, after=4, line=1.15)

    for style_name in ("List Bullet", "List Number"):
        set_font(styles[style_name], name="Calibri", size=11, color="000000")
        set_spacing(styles[style_name], before=0, after=4, line=1.15)
        styles[style_name].paragraph_format.left_indent = Inches(0.5)
        styles[style_name].paragraph_format.first_line_indent = Inches(-0.25)

    code_style = styles.add_style("Manuscript Code", 1)
    set_font(code_style, name="Consolas", size=9, color="111827")
    set_spacing(code_style, before=2, after=2, line=1.0)
    code_style.paragraph_format.left_indent = Inches(0.25)
    code_style.paragraph_format.right_indent = Inches(0.15)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("RailRL manuscript draft v0")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("666666")


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_inline(paragraph, text: str) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def flush_paragraph(doc: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    text = " ".join(s.strip() for s in buffer).strip()
    if text:
        p = doc.add_paragraph()
        add_inline(p, text)
    buffer.clear()


def is_table_separator(line: str) -> bool:
    parts = [p.strip() for p in line.strip().strip("|").split("|")]
    return bool(parts) and all(re.fullmatch(r":?-{3,}:?", p or "") for p in parts)


def parse_table_row(line: str) -> list[str]:
    return [p.strip() for p in line.strip().strip("|").split("|")]


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    rows = [r + [""] * (ncols - len(r)) for r in rows]
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    table.autofit = False
    usable_dxa = 9360
    widths = [usable_dxa // ncols] * ncols
    widths[-1] += usable_dxa - sum(widths)
    for j, text in enumerate(rows[0]):
        cell = table.rows[0].cells[j]
        set_cell_width(cell, widths[j])
        shade_cell(cell)
        p = cell.paragraphs[0]
        add_inline(p, text)
        for run in p.runs:
            run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for j, text in enumerate(row):
            set_cell_width(cells[j], widths[j])
            add_inline(cells[j].paragraphs[0], text)


def build_docx(src: Path = SRC, out: Path = OUT) -> None:
    text = src.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)

    paragraph_buffer: list[str] = []
    in_code = False

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.strip().startswith("```"):
            flush_paragraph(doc, paragraph_buffer)
            in_code = not in_code
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph(style="Manuscript Code")
            p.add_run(line)
            add_shading(p)
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            flush_paragraph(doc, paragraph_buffer)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            flush_paragraph(doc, paragraph_buffer)
            table_rows = [parse_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_markdown_table(doc, table_rows)
            continue

        if stripped == "---":
            flush_paragraph(doc, paragraph_buffer)
            p = doc.add_paragraph()
            add_bottom_border(p)
            i += 1
            continue

        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            flush_paragraph(doc, paragraph_buffer)
            level = len(m.group(1))
            title = m.group(2).strip()
            if level == 1:
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_inline(p, title)
            else:
                p = doc.add_paragraph(style=f"Heading {level - 1}")
                add_inline(p, title)
            i += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph(doc, paragraph_buffer)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:].strip())
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            flush_paragraph(doc, paragraph_buffer)
            item = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            add_inline(p, item)
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph(doc, paragraph_buffer)
    doc.save(out)
    print(out)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--src", type=Path, default=SRC)
    parser.add_argument("--out", type=Path, default=OUT)
    args = parser.parse_args()
    build_docx(args.src, args.out)
