"""Generate the comprehensive project handoff docx for RailRL v2.

Writes to: docs/PROJECT_HANDOFF.docx

This is a from-scratch handoff document compiled from:
- RailRL_v1 phase2_feature_spec.md (v2.2)
- Research_Proposal_Derby_RL_v3.docx
- All audit findings from outputs/analyses/*.json
- Reward calibration outputs
- Conversation history about v2 redesign
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# helpers
# ============================================================

def set_cjk_font(run, font_name="Microsoft YaHei"):
    """Set CJK-friendly font on a run."""
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfont = rpr.find(qn('w:rFonts'))
    if rfont is None:
        rfont = OxmlElement('w:rFonts')
        rpr.append(rfont)
    rfont.set(qn('w:eastAsia'), font_name)
    rfont.set(qn('w:ascii'), font_name)
    rfont.set(qn('w:hAnsi'), font_name)


def add_heading(doc, text, level):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.bold = True
    set_cjk_font(run, "Microsoft YaHei")
    if level == 0:
        run.font.size = Pt(20)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    else:
        run.font.size = Pt(11)
    return h


def add_para(doc, text, bold=False, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    if isinstance(text, list):
        for chunk, b, i in text:
            run = p.add_run(chunk)
            run.bold = b
            run.italic = i
            run.font.size = Pt(size)
            set_cjk_font(run, "Microsoft YaHei")
    else:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        set_cjk_font(run, "Microsoft YaHei")
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.7 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    set_cjk_font(run, "Microsoft YaHei")
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    # background shading
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def add_table(doc, header, rows, col_widths_cm=None):
    """Add a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # widths
    if col_widths_cm:
        for j, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[j].width = Cm(w)

    # header
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(header):
        hdr_cells[j].text = ""
        p = hdr_cells[j].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        set_cjk_font(run, "Microsoft YaHei")
        # header shading
        tcPr = hdr_cells[j]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D5E8F0')
        tcPr.append(shd)

    # data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            set_cjk_font(run, "Microsoft YaHei")

    # space after
    doc.add_paragraph()
    return table


def add_callout(doc, title, body, color="FFF4E1"):
    """Highlighted callout box (single-cell table)."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)
    cell.text = ""
    p = cell.paragraphs[0]
    if title:
        run = p.add_run(title + "\n")
        run.bold = True
        run.font.size = Pt(11)
        set_cjk_font(run, "Microsoft YaHei")
    run = p.add_run(body)
    run.font.size = Pt(10)
    set_cjk_font(run, "Microsoft YaHei")
    doc.add_paragraph()
    return t


# ============================================================
# build document
# ============================================================

doc = Document()

# Default style
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# Page setup
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# ============================================================
# TITLE PAGE
# ============================================================
add_heading(doc, "RailRL v2 项目说明文档", 0)
add_para(doc, "Derby 工作站信号员决策的端到端可解释强化学习系统", bold=True, size=13)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("目标期刊：ESWA（Expert Systems with Applications）")
set_cjk_font(run); run.font.size = Pt(11); run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"维护者：Hao  ·  最后更新：2026 年 5 月 18 日")
set_cjk_font(run); run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()

add_callout(doc, "本文档的用途",
"本文档是 RailRL v2 项目的完整交接文档。当你开启一个新的对话窗口讨论本项目时，"
"把这份文档作为第一篇阅读资料给到 AI 助手——它包含了项目问题定义、信号员的真实工作流程、"
"数据契约、为什么选择强化学习、当前架构设计、v1 失败教训、当前进度和后续计划等所有必要上下文。"
"目标是：新对话不再需要从头解释整个项目背景。")

doc.add_paragraph()

# Table of contents
add_heading(doc, "目录", 1)
toc_items = [
    "第 1 章   项目背景与定位",
    "第 2 章   业务领域：信号员的真实工作",
    "第 3 章   三类决策详解（核心）",
    "第 4 章   数据资源详细说明",
    "第 5 章   关键实证发现（来自 v1 数据审计）",
    "第 6 章   为什么强化学习适合这个项目",
    "第 7 章   核心架构设计（v2）",
    "第 8 章   信息泄露契约（critical）",
    "第 9 章   五层可解释性",
    "第 10 章  评估框架",
    "第 11 章  v1 失败教训（必须记住）",
    "第 12 章  项目进度与计划",
    "第 13 章  奖励详解（4 维 r 怎么算）   ← 2026-05-19 新增",
    "第 14 章  Leak 审计扩展（强制 checklist）  ← 2026-05-19 新增",
    "第 15 章  实施路线图（主模型优先）   ← 2026-05-19 新增",
    "附录 A    常用术语表",
    "附录 B    关键文件路径",
    "附录 C    关键数字一览",
]
for item in toc_items:
    add_para(doc, item, size=10.5)

doc.add_paragraph()
add_callout(doc, "更新日志 (CHANGELOG)",
"  • 2026-05-18 v1.0：初版文档（12 章 + 3 附录）\n"
"  • 2026-05-19 v1.1：新增第 13/14/15 章 + 更新 Ch 2.6 / 7.4\n"
"    - 第 13 章：4 维 reward 的详细计算公式（r_delay / r_throughput / r_headway / r_wait）\n"
"    - 第 14 章：Leak 审计 checklist（由 reward 中间量引申的禁止字段 + assert_no_leak）\n"
"    - 第 15 章：实施路线图（用户决策：主模型优先 → baselines 后置）\n"
"    - Ch 2.6 priority 修正：FCFS 易模仿但可改进，priority head 为 improvement 而保留\n"
"    - Ch 7.4 Q 网络：明确 Derby_info 物理特征进入 route_emb",
color="E1EFFF")

doc.add_page_break()

# ============================================================
# Chapter 1 — Background
# ============================================================
add_heading(doc, "第 1 章  项目背景与定位", 1)

add_heading(doc, "1.1 一句话定位", 2)
add_para(doc,
"用 14 个月真实的 Network Rail 开放数据，从 UK Derby 火车站工作站的信号员历史按钮记录中，"
"端到端学习一个能够联合做出"
"（i）进路设置、（ii）信号办理时机、（iii）资源冲突优先级"
"三类决策的可解释离线强化学习系统。"
)

add_heading(doc, "1.2 目标期刊：ESWA（Expert Systems with Applications）", 2)
add_para(doc,
"ESWA 是 Elsevier 旗下的应用型 AI / 专家系统期刊，特别看重："
)
add_bullet(doc, "完整的、可部署的端到端系统，而不仅是算法 benchmark")
add_bullet(doc, "真实数据 + 真实业务背景")
add_bullet(doc, "充分的可解释性 / 可问责性")
add_bullet(doc, "详尽、可复现的实验评估")

add_para(doc, "这与我们项目的特点高度匹配——14 个月真实信号员数据 + 完整的数据采集 pipeline + 五层可解释决策 + Replicate-AND-Improve 评估框架。")

add_heading(doc, "1.3 论文三大贡献（围绕 ESWA 章节展开）", 2)

add_table(doc,
    header=["章节", "贡献", "对应代码模块"],
    rows=[
        ["§3 + §4", "数据采集 + 工程化 pipeline：从 Network Rail 开放数据流到 MDP 元组", "src/railrl/data/"],
        ["§5", "端到端 RL 框架：联合学习路径 + 时机 + 优先级三类决策", "src/railrl/mdp/, encoders/, policies/, algorithms/"],
        ["§6–§8", "五层可解释决策 + Replicate-AND-Improve 评估 + Selective Override 部署规则", "src/railrl/xai/, eval/, deploy/"],
    ],
    col_widths_cm=[2.0, 9.0, 5.5],
)

add_heading(doc, "1.4 项目核心叙事：Replicate-AND-Improve", 2)
add_para(doc,
"v3 提案中确立的关键叙事转变："
)
add_callout(doc, "核心叙事",
"不再把信号员的行为当作金标准。14 个月的演示数据存在质量异质性"
"（换班、疲劳、突发交通、信息延迟等都会让某些 PR 操作次优）。"
"目标是：当信号员正确时复现他，当信号员次优时提出更好的替代方案，并对两者都给出可解释的依据；"
"只有当 L3 反事实仿真证明替代方案在系统层面更好 AND L4 手册规则认为合规 AND L2 决策解释可信时，"
"才打断人类操作员。")

add_para(doc,
"这种 selective override 的部署规则比"
"\"始终跟随信号员\""
"或"
"\"始终用模型预测\""
"都要诚实和可用——这是项目想要交付的真正价值。"
)

doc.add_page_break()

# ============================================================
# Chapter 2 — Signaller's Real Work
# ============================================================
add_heading(doc, "第 2 章  业务领域：信号员的真实工作", 1)

add_heading(doc, "2.1 什么是 route setting（进路设置）", 2)
add_para(doc,
"在英国铁路网络上，"
"route setting"
"是信号员通过控制台按按钮，"
"为即将通过某个 signal（信号机）的列车锁定一条物理路径（route）的行为。"
"每按一次按钮，interlocking 系统会锁定一组 track sections（TCs，轨道电路），"
"使信号机变绿、点机器（points）摆到正确位置，列车才能安全通过。"
)

add_heading(doc, "2.2 Derby 工作站的物理布局", 2)
add_table(doc,
    header=["资产类型", "数量", "说明"],
    rows=[
        ["Signals", "123", "信号机；每条 route 都从某个 signal 出发"],
        ["Track sections", "249", "轨道电路；列车占用某 TC 时会产生 TD Track 事件"],
        ["Routes", "277", "命名路径；route_to_tc_all.csv 给出每条 route 经过的有序 TC 列表"],
        ["Aux connections", "156", "辅助连接（非 panel 内部）"],
        ["Active trains 平均", "1–15", "繁忙时刻最多 15 辆同时在 approach 域内"],
    ],
    col_widths_cm=[3.0, 2.0, 11.0],
)

add_para(doc, "Derby 按 line-of-route 划分为 5 个 prefix 子区域：", bold=True)
add_table(doc,
    header=["Prefix", "全称", "在数据中的 PR 数（14 个月）"],
    rows=[
        ["DW", "Derby ↔ Wichnor Junction（南向走廊）", "120,894"],
        ["TD", "Trent Junction ↔ Derby（东侧 approach + 站内）", "145,165"],
        ["DC", "Derby ↔ Chesterfield（北向走廊）", "194,796"],
        ["EC", "Etches Park ↔ Chaddesden（depot / sidings）", "27,343"],
        ["DY", "Duffield ↔ 北端（re-controlled 段）", "58,220"],
    ],
    col_widths_cm=[2.0, 9.0, 5.0],
)
add_para(doc, "总计 546,418 条 Panel_Request 决策事件（横跨 2023-02-28 至 2024-04-25）。", italic=True)

add_heading(doc, "2.3 信号员在每一刻的真实操作流程", 2)
add_para(doc, "信号员盯着一个面板（panel），面板上显示：")
add_bullet(doc, "所有 TC 的占用状态（哪些被列车占了）")
add_bullet(doc, "所有 signal 的当前 aspect（红/黄/绿）")
add_bullet(doc, "当前已锁定的 routes")
add_bullet(doc, "列车的 4 字符 headcode（如 1S49）和当前 berth（停车位）")
add_bullet(doc, "时刻表（gbtt_timestamp 计划时刻、planned_platform 计划站台）")

add_para(doc, "他每隔几秒看一次面板，然后做一个动作——")
add_callout(doc, "信号员的动作（每一时刻）",
"要么按某个按钮（为某辆车设某条进路），要么什么都不做（等几秒再看）。\n"
"这两种情况之间的选择，就是 (i) 时机决策。\n"
"按按钮时选哪辆车的哪条路径，就是 (ii) 优先级 + (iii) 进路决策。\n"
"三类决策同时蕴含在一个物理动作里。",
color="E1EFFF")

add_heading(doc, "2.4 一个具体的工作场景（按图解释）", 2)
add_para(doc, "信号员面前的 panel 显示了一辆 1S49（class 1 = express passenger）已经通过 signal DC5079——")

add_para(doc, "1. 信号员知道：1S49 接下来必然要走到 signal 5087 附近。", bold=True)
add_para(doc, "2. 5087 出发的 routes 有 2 个候选：前往 5093 的，或前往 5095 的。", bold=True)
add_para(doc, "3. 信号员要决定的不是\"走 5093 还是 5095\"（schedule 大多数情况已经决定了），而是：", bold=True)
add_bullet(doc, "什么时候按 5087 → 5093 的按钮？")
add_bullet(doc, "如果 2A28（前面一辆车）还在 5093 路径的某个 TC 上，是再等几秒还是 advance routing？")
add_bullet(doc, "如果另一辆同时到达的 8M28 也需要 5087 区域的资源，谁先按？")

add_callout(doc, "关键洞察",
"在 90% 的决策中，路径（route）的选择是由 schedule + 图论 + 列车类别共同决定的"
"——\"伪决策\"。"
"真正考验信号员经验的是剩下 10% 的：advance routing、call-on、platform 临时改派、"
"priority swap、recovery routing 等。"
"这 10% 才是论文真正要捕捉的价值所在。",
color="FFF4E1")

add_heading(doc, "2.5 时机决策：太早 vs 太晚的代价", 2)
add_para(doc, "信号员对\"什么时候按按钮\"非常敏感，因为太早和太晚都有代价：")

add_table(doc,
    header=["时机", "后果", "影响"],
    rows=[
        ["太早按", "前一辆车还没出清，造成 route 冲突 / interlocking 锁定", "他车被卡住、headway 不足、潜在安全风险"],
        ["太晚按", "本车看到的 signal 是红，被迫降速", "列车 dwell 增加、整体 delay 累积、影响下游"],
        ["刚好", "前车刚出清，本车不减速顺利通过", "delay 最小化、throughput 最大化"],
    ],
    col_widths_cm=[2.5, 6.0, 7.5],
)

add_para(doc, "这是一个典型的\"最优停时\"（optimal stopping）问题——而且没有公式，只能从信号员的历史决策里学。")

add_heading(doc, "2.6 优先级决策：多车竞争场景", 2)
add_para(doc, "Derby 在繁忙时段会同时有 10–15 辆活跃列车，其中若干会同时需要相邻 signal 的资源。例子：")
add_bullet(doc, "Express 1S49 和货车 6M91 都从北边进来 → 一般客车优先")
add_bullet(doc, "两辆 Express 同时到 platform 3 区域 → 先到先得 / 看具体进站方向")
add_bullet(doc, "ECS（空车调度）5K23 与正点客车 1S49 → 让 1S49 先走，5K23 让")

add_para(doc, "信号员的优先级判断不是公式，但是非常一致——同一类场景的处理模式高度可预测。这正是我们的模型要去恢复的隐性偏好。")

add_callout(doc, "重要修正（2026-05-19 讨论）：priority 是「易模仿、可改进」",
"领域参考文献观察：在 Derby 单 panel + 5 秒并发窗口下，FCFS（先到先服务）启发式能预测信号员的 "
"priority 决策达到 Kendall τ ≈ 0.998，几乎完美。这意味着：\n\n"
"  · 信号员的 priority 决策几乎可被 FCFS 完美模仿（imitation 是 trivial 的）\n"
"  · 但 FCFS 能预测信号员 ≠ FCFS 就是 4 维 reward 下的最优策略\n\n"
"FCFS 可能只是信号员默认使用的低认知负载启发式；从 reward 角度（delay/throughput/headway/wait）"
"看，CQL 可能找到非 FCFS 的更优 priority 排序——比如让晚点客车插货车之前。\n\n"
"因此 priority head 在 v2 仍然保留，但定位修正为：\n"
"  · 不为了「预测 priority 准确」（FCFS 已经准）\n"
"  · 为了「surface CQL 在保守 support 内的 policy improvement opportunity」\n"
"  · 评估两个指标分开报：(a) imitation Kendall τ（预期 ≈ 0.998）+ "
"(b) counterfactual reward delta（CQL 选 vs 信号员/FCFS 选，L3 仿真对比 reward 改善）\n\n"
"这正是 Replicate-AND-Improve 叙事在 priority 维度的具体体现。"
"⚠ 上述 0.998 数字来自 RailMind 项目的实测，仅作领域参考；v2 需在自己的数据上独立复现验证。",
color="FFE9E9")

add_heading(doc, "2.7 信号员的高价值经验（必须捕捉）", 2)
add_para(doc, "这 6 类决策占总决策的 10% 左右，但它们是信号员存在的意义：")

add_table(doc,
    header=["类别", "占比", "信号员隐性 know-how"],
    rows=[
        ["Advance routing（提前布线）", "~3%", "前车马上要出清某 TC，提前几秒按可省 30 秒 dwell——但再早 2 秒就 conflict"],
        ["Call-on（呼叫路径）", "~1.5%", "1S49 急着进 3 站台但 2S88 还在 dwell——开 C-class 让 1S49 慢速跟进"],
        ["Platform 改派", "~1.5%", "1S49 原定 platform 3 但被 8M28 占用——改 platform 1，提前 set 5087 那条而不是 5093"],
        ["Priority swap（优先级反转）", "~2%", "6M91 货车其实可以让 1S49 客车先过——晚 set freight 路径，先 set passenger 路径"],
        ["Recovery routing", "~1%", "1S49 已经晚 5 分钟——用 main 而非 calling-on，省每一秒"],
        ["Emergency / 非标", "~1%", "对 343R（非标 train_id）做特殊处理；TPWS 故障；点机器故障"],
    ],
    col_widths_cm=[3.5, 1.5, 11.0],
)

add_callout(doc, "评估的真正靶心",
"论文 §VII 的主结果不在 Overall accuracy（会被 90% trivial 主导，掩盖模型真正能力），"
"而在上面这 6 个特殊情况列上每一列的精度——HG-DT-CQL 在这 10% 决策上的提升才是真正可发表的发现。",
color="FFE9E9")

add_heading(doc, "2.8 时刻表的不确定性", 2)
add_para(doc,
"时刻表（schedule）告诉信号员每辆车的"
"\"计划到达时间\""
"和"
"\"计划停靠的 platform\""
"，但它不是硬约束——")
add_bullet(doc, "列车可能临时被改派到另一个空 platform")
add_bullet(doc, "列车可能晚点 / 早点，导致原定 platform 还被占用")
add_bullet(doc, "列车 class（headcode 第一字符）会影响停靠/路过的选择")

add_para(doc, "所以模型不能把 planned_platform 当作 hard mask 去 filter 候选 routes，否则会自动剔除\"信号员真实做出的 platform 改派决策\"——这正是 §V 想抓的高价值经验。")

doc.add_page_break()

# ============================================================
# Chapter 3 — Three Decisions
# ============================================================
add_heading(doc, "第 3 章  三类决策详解（核心）", 1)

add_heading(doc, "3.1 三类决策的形式定义", 2)

add_table(doc,
    header=["决策类型", "信号员真实做的事", "数据中的体现"],
    rows=[
        ["(i) 进路设置 Route setting",
         "选择为当前列车设哪条 route_id",
         "decision_events.parquet 的 chosen_route_id 列"],
        ["(ii) 信号办理时机 Timing",
         "选择什么时候按下 PR 按钮",
         "approach 触发时刻到 PR 时刻的时间差"],
        ["(iii) 资源冲突优先级 Priority",
         "多车同时活跃时选择服务哪辆",
         "同窗口内多车竞争场景下信号员实际服务的列车"],
    ],
    col_widths_cm=[4.0, 6.5, 5.5],
)

add_heading(doc, "3.2 关键洞察：三类决策可统一为一个结构化动作", 2)

add_para(doc, "在任意时刻 t，可选动作集合是动态、结构化的：", bold=True)

add_code(doc,
"A_t = { wait }                                  ← 1 个 wait 选项\n"
"      ∪\n"
"     { (列车_1, 路径_a),\n"
"       (列车_1, 路径_b),\n"
"       (列车_2, 路径_c),\n"
"       (列车_2, 路径_d),\n"
"       ...                                       ← 通常 3-30 个 (列车,路径) 二元组\n"
"     }\n"
"\n"
"  繁忙时刻 |A_t| ≈ 20-30\n"
"  空闲时刻 |A_t| ≈ 3-8\n"
"  完全空闲 |A_t| = 1（只能 wait）"
)

add_para(doc, "信号员选了 (列车_2, 路径_d) 一个动作，意味着三件事同时发生：")
add_bullet(doc, "优先级：选了列车 2，没选列车 1 和列车 3 → 列车 2 优先级最高")
add_bullet(doc, "时机：在这个 tick 选了非-wait → 现在就办，没再等")
add_bullet(doc, "路径：在列车 2 的 3-5 条可达路径里挑了 d → 路径选了 d")

add_callout(doc, "为什么这样建模是对的",
"信号员物理上做的就是\"在某一时刻按某个按钮\"——一个按钮蕴含三类决策。"
"v1 把它强行拆成\"set vs wait\"二分类，路径信息和优先级信息全部丢失，"
"导致 91% 准确率几乎全部由 focal_signal 的 one-hot 编码给出。"
"v2 用结构化动作空间避免了这个问题，三类决策自然耦合在一个动作里。",
color="E1EFFF")

add_heading(doc, "3.3 候选动作集合不是 277-class 分类", 2)
add_para(doc,
"虽然 Derby 有 277 条命名 routes，"
"但在任意时刻 t，对任意焦点列车，物理可达且语义合理的路径只有 3-10 条。"
"逐层缩减如下："
)

add_code(doc,
"全部 named routes                                            277\n"
"  └─ 从列车当前 TC 物理可达（图 BFS）                          5 - 15\n"
"       └─ 在列车 headcode_class 历史上出现过                     3 - 8\n"
"            └─ 与列车 planned_platform 一致                       2 - 4\n"
"                 └─ 当前 traffic state 下 interlocking 允许       1 - 3"
)

add_para(doc, "v1 实证支撑（来自 outputs/analyses/route_class_summary.json）：")
add_bullet(doc, "EC-prefix 少数派路径 95%+ 是 ECS（class 5）—— 路径与列车类别强相关")
add_bullet(doc, "Call-on (C-class) 路径 92% 是 Express（class 1）—— 特定语义场景")
add_bullet(doc, "Derby 每个 signal 平均 2.3 条 outbound routes（277 ÷ 123 signals）")

doc.add_page_break()

# ============================================================
# Chapter 4 — Data
# ============================================================
add_heading(doc, "第 4 章  数据资源详细说明", 1)

add_heading(doc, "4.1 三个原始数据源", 2)

add_table(doc,
    header=["文件", "大小", "内容", "用途"],
    rows=[
        ["TD_data.csv", "713 MB", "TD S-class + C-class 事件，11.91M 行，含 SOP 解码后的 5 类事件", "状态 + 动作标签源头"],
        ["Movements.csv", "49 MB", "TRUST 实际运行事件，247k 行，含 gbtt / actual / variation timestamps", "schedule + reward 信号源"],
        ["route_to_tc_all.csv", "32 KB", "Derby 路径基础设施，447 行，每行一条 route 经过的有序 TC 列表", "异构图静态骨架"],
    ],
    col_widths_cm=[3.5, 1.5, 7.5, 4.5],
)

add_heading(doc, "4.2 TD S-class 事件的 5 个类别（来自 SOP 解码）", 2)
add_para(doc, "原始 S-class 报文的每个 byte 经 Algorithm 3.1（论文 Chapter 3）解码后落到 5 类：")

add_table(doc,
    header=["事件类型", "含义", "在 MDP 中的角色"],
    rows=[
        ["Track", "TC 占用 / 释放", "状态：track 节点 occupied/cleared"],
        ["Signal", "信号机 aspect 变化", "状态：signal 节点 aspect"],
        ["Route", "Route 锁定 / 释放", "状态：route 节点 locked/cleared"],
        ["Panel_Request", "信号员按了 PR 按钮（state=1）", "动作：chosen_route_id"],
        ["TRTS", "Train Ready To Start（站台 ready）", "状态特征：trts_pressed flag"],
    ],
    col_widths_cm=[3.0, 5.0, 8.0],
)

add_para(doc, "C-class 事件（CA / CB / CC）追踪 berth 移动，用于：")
add_bullet(doc, "train 节点的 current_berth 特征")
add_bullet(doc, "推断列车 trajectory（recent_tcs）")
add_bullet(doc, "TRUST id ↔ TD trainid_filled 的 pass disambiguation")

add_heading(doc, "4.3 参考数据（v2 仓库内 data/reference/）", 2)
add_table(doc,
    header=["文件", "用途"],
    rows=[
        ["platform_end_signals.csv", "每个 platform 末端的 signal 映射"],
        ["platform_tc_map.csv", "TC ↔ platform_id 双向映射 + sub_section（A/middle/B）"],
        ["Derby_info.csv", "Derby SOP 资产 ID → name 映射"],
        ["derby_info_mapping.csv", "asset_idx (0-672) ↔ asset_name 双向索引"],
        ["TRT1.DY2_2.SOP", "原始 SOP 字典文件"],
        ["derby_all.png", "Derby 平面示意图，用于 L1 attention 可视化"],
    ],
    col_widths_cm=[5.5, 10.0],
)

add_heading(doc, "4.4 领域文档（v2 仓库内 data/domain/）", 2)
add_table(doc,
    header=["文件", "用途"],
    rows=[
        ["Training_Plan_2022.docx", "Derby Workstation Training Plan：preferred routes、call-on 规则、§3 + §5"],
        ["Signalling_Nomenclature.pdf", "UK 信号系统标准命名规范"],
        ["headcode.pdf", "Train headcode 4 字符编码规则"],
        ["S-class.pdf", "S-class 事件 SOP 字节级解码规范"],
    ],
    col_widths_cm=[5.5, 10.0],
)

add_heading(doc, "4.5 已生成的 Phase 2 工件（v2 仓库内 outputs/）", 2)
add_table(doc,
    header=["目录", "关键文件", "规模"],
    rows=[
        ["inventory/", "td_inventory.json, movements_inventory.json", "<100 KB"],
        ["decisions/", "decision_events.parquet（含 chosen_route_id）", "5.3 MB / 546k 行"],
        ["infrastructure/", "routes_clean / tracks_inventory / signals_inventory parquet", "~50 KB"],
        ["static_graph/", "4 节点类型 × 6 边类型的全部 parquet + summary.json", "~90 KB"],
        ["event_stream/", "event_tokens.parquet（K=256 token 流）", "42 MB"],
        ["rewards/", "calibration.json + pr_outcomes.parquet + decision_rewards.parquet", "~14 MB"],
        ["analyses/", "3 个实证分析 parquet + 3 个 summary.json", "~3 MB"],
        ["cache/", "td_data.parquet（TD 解析缓存）", "90 MB"],
    ],
    col_widths_cm=[3.5, 8.0, 4.0],
)

add_heading(doc, "4.6 数据使用的核心约束：时序因果契约", 2)
add_callout(doc, "Temporal Causality Contract（来自 phase2_feature_spec.md §B.0）",
"决策点 t 的 state 只能用 time ≤ t 的事件构造（因果可见）。\n\n"
"  • 调度信息（gbtt_timestamp, planned_timestamp）→ ✅ state 可用（信号员有时刻表）\n"
"  • 实际信息（actual_timestamp 在 t' > t）→ ❌ state 禁用（未来未发生）\n"
"  • TD 事件 time > t 的 → ❌ state 禁用\n"
"  • timetable_variation at t' > t → ❌ state 禁用，但 ✅ reward 可用（return 是 hindsight）",
color="FFE9E9")

doc.add_page_break()

# ============================================================
# Chapter 5 — Empirical findings
# ============================================================
add_heading(doc, "第 5 章  关键实证发现（来自 v1 数据审计）", 1)
add_para(doc, "v1 数据工程阶段做了 8 条经过验证的数据审计发现，这些都是设计决策的依据，论文 §3 中是 ESWA 审稿人最爱的\"数据驱动决策\"叙事素材。", italic=True)

add_heading(doc, "5.1 Conflict 经验分析", 2)
add_callout(doc, "Finding 1：97.6% 的\"conflict edges\"是假阳性",
"对 546,111 个 PR 做 audit，发现 95.22% 的 PR 触发时 conflict mask 内的 TC 完全没被占用；"
"剩下 2.19% 是被他车占用，但其中绝大多数是 queue routing（前车马上要出清的合法 advance routing）。\n\n"
"决策：v2 把 conflict 维度完全从 reward 和 action mask 移除。"
"interlocking 系统会处理物理安全，模型只学信号员的按钮时机。"
"占用信息保留为 state 特征。")

add_heading(doc, "5.2 L1 教科书规则的违反", 2)
add_callout(doc, "Finding 2：3.01% 的真实 PR 违反教科书 L1 first-TC-clear 规则",
"信号员明确使用 advance routing（前车未完全出清就提前布线）—— 这是 expertise 的体现，"
"不是\"违规\"。\n\n"
"决策：L1 hard mask 默认关闭；advance routing 作为 special-case flag 显式喂给 Q 网络。")

add_heading(doc, "5.3 路径与列车类别的强相关", 2)
add_callout(doc, "Finding 3：少数派路径高度类别专门化",
"  • EC-5475 B(M)：5.6% signal traffic，98.3% 是 ECS（class 5）\n"
"  • DC-5076 A(M)：3.9% signal traffic，98.1% 是 ECS\n"
"  • DW-5306 B(C) Call-on：9.7%，91.9% 是 Express（class 1）\n"
"  • TD-5045 A(C) Call-on：1.6%，91.8% 是 Express\n\n"
"决策：route_class × headcode_class 的强相关让 GNN 通过 attention 自然学到，不需要额外特征工程。")

add_heading(doc, "5.4 Headcode 7/8 处理", 2)
add_callout(doc, "Finding 4：headcode class 7 和 8 是\"non-standard\"，合并处理",
"  • Class 7：442 条 PR\n"
"  • Class 8：28 条 PR\n\n"
"决策：将 7 和 8 都映射到 \"other\" 类别。")

add_heading(doc, "5.5 非标准 train_id", 2)
add_callout(doc, "Finding 5：1.04% 的 train_id 是非标准 4-char 格式（如 \"343R\"）",
"这些 ID 在 depot / sidings / shunt 场景下集中出现：\n"
"  • EC-prefix 路径中过度出现 3.5× 倍\n"
"  • Shunt-class 路径中过度出现 6.7× 倍\n\n"
"决策：保留这些行（不丢弃），hc_class_digit 设为 'non_standard' 类别，"
"作为 f_unusual_id flag 喂给 Q 网络。")

add_heading(doc, "5.6 Reward 校准的经验阈值", 2)
add_para(doc, "P2.4 阶段从 14 个月全量数据用 percentile 校准的 4 个关键阈值：", bold=True)

add_table(doc,
    header=["参数", "经验值", "百分位 / 含义"],
    rows=[
        ["H_min（最小可接受 headway）", "147.0 s", "P5 of 3.28M pair-wise headways"],
        ["d-gate 0.5 boundary", "6 hops", "P50 of 24,870 sampled decisions"],
        ["d-gate 0.1 boundary", "16 hops", "P90 of sampled decisions"],
        ["Reward observation window", "4201.9 s", "P99 of 42,806 TIPLOC lags"],
    ],
    col_widths_cm=[5.0, 3.0, 7.5],
)

add_heading(doc, "5.7 Per-signal majority baseline 揭示的任务结构", 2)
add_callout(doc, "Finding 6：per-signal-majority = 91% accuracy（v1 binary 任务下）",
"这是\"trivial baseline\"——只要知道是哪个 signal，按经验最常见做法选 set/wait 就能 91%。\n\n"
"含义：71% 的决策落在\"行为模式极度一致\"的 signal 上（如 5037/5040 几乎总 set，5090/5096 几乎总 wait）。\n\n"
"决策：必须用结构化动作空间（277-class，不是 2-class），让 trivial baseline 降到 30-50%。")

add_heading(doc, "5.8 Reward 的分布质量", 2)
add_para(doc, "在 2.64M 决策、82,429 episodes 上的健康检查（outputs/rewards/health/health_summary.md）：")
add_bullet(doc, "r_total 均值 +0.255、标准差 0.675、范围 [-30.30, +30.50]")
add_bullet(doc, "Per-episode return 均值 +2.25，88.4% episodes 正回报")
add_bullet(doc, "Weight 排序稳健性：conservative vs default Spearman = 0.908")
add_bullet(doc, "Movements proxy 验证：Spearman(r_total, delay_reduction) = +0.29")

doc.add_page_break()

# ============================================================
# Chapter 6 — Why RL
# ============================================================
add_heading(doc, "第 6 章  为什么强化学习适合这个项目", 1)

add_heading(doc, "6.1 监督学习（imitation / BC）的瓶颈", 2)
add_para(doc, "如果只做行为克隆（BC），论文的天花板就是信号员本身的水平。两个问题：")
add_bullet(doc, "演示数据质量异质：14 个月里换班、疲劳、压力、信息延迟都会让某些 PR 操作次优")
add_bullet(doc, "BC 只学 π，不学 Q——无法判断\"信号员当时的选择是否最优\"")

add_heading(doc, "6.2 仿真训练（model-based RL）的不可行", 2)
add_bullet(doc, "我们没有 ground-truth 环境模型；只有 14 个月的历史轨迹")
add_bullet(doc, "想要的恰恰是\"恢复信号员实际偏好\"——仿真环境的 reward 函数本身就是要学的目标")
add_bullet(doc, "L3 反事实仿真器（P2.6）是评估工具，不是训练工具")

add_heading(doc, "6.3 离线 RL（offline RL）的契合点", 2)
add_callout(doc, "为什么 CQL / IQL 正好适合本项目",
"  1. 14 个月演示数据是完美的 offline RL buffer（546k+ PR + reward）\n"
"  2. Conservative 算法（CQL）在异质质量演示上有理论保证：\n"
"     - 同一 state 多个演示动作时，倾向 in-support 且高 Q 的动作\n"
"     - 同质演示时，等价 BC（保守性自动失活）\n"
"     - 这就是\"replicate when right, improve when sub-optimal\"的机制\n"
"  3. 训练目标是 Q(s, a)，自带\"哪些动作更好\"的可解释信息\n"
"  4. 不需要与环境交互——真实环境是 UK 铁路网，没法 rollout",
color="E1EFFF")

add_heading(doc, "6.4 端到端联合三类决策的需要", 2)
add_para(doc,
"v1 把任务拆成 binary（set / wait）+ 单独的 route head + ...，"
"结果三类决策被错误地解耦：")
add_bullet(doc, "binary 任务自动丢失路径信息（277 → 2 类）")
add_bullet(doc, "优先级决策无处建模（多车场景被压成单一决策）")
add_bullet(doc, "时机决策也丢失（set/wait 不区分\"再等 10 秒\"和\"再等 60 秒\"）")

add_para(doc,
"v2 用结构化动作 {wait} ∪ {(train, route)} 让三类决策"
"自然耦合在一个 Q 函数里——这是 RL 框架的天然优势：")
add_bullet(doc, "单一 Q(s, a) 同时给出\"哪辆车（优先级）+ 哪条路径（路径）\"的最优选择")
add_bullet(doc, "在时间序列上，wait vs 非-wait 的选择构成时机决策")
add_bullet(doc, "无需 hand-engineered 多 head 加权拼接")

add_heading(doc, "6.5 与 Replicate-AND-Improve 叙事的契合", 2)
add_para(doc,
"CQL 的保守性正好支撑论文的核心叙事——")
add_bullet(doc, "在数据 support 范围内，CQL 不偏离信号员太远（replicate）")
add_bullet(doc, "在 support 内有更优的高 Q 动作时，CQL 会选它（improve）")
add_bullet(doc, "L2 的 Q-gap 给出可解释的\"为什么 model 偏离信号员\"")
add_bullet(doc, "Selective override 用 L2/L3/L4 三重 gate 决定是否真的 override")

doc.add_page_break()

# ============================================================
# Chapter 7 — Architecture
# ============================================================
add_heading(doc, "第 7 章  核心架构设计（v2）", 1)

add_heading(doc, "7.1 MDP 形式化", 2)
add_table(doc,
    header=["元素", "v2 定义"],
    rows=[
        ["State s_t", "异构图 snapshot（4 节点类型 × 6 边类型）+ K=256 事件序列 + 多车上下文 + schedule outlook，无 focal_signal"],
        ["Action a_t", "{wait} ∪ {(train_id, route_id)}，候选 mask 从图 BFS + direction + 已 set 历史 + planned_platform 推得"],
        ["Reward r_t", "P2.4 校准的 4 维：r_delay + r_throughput + r_headway + r_wait（默认权重 1.0/0.5/1.0/0.3）"],
        ["Episode", "每辆车一段：从该 train 首次进入 approach horizon 到离开 Derby（或 pass_id 结束）"],
        ["Discount γ", "0.95（约 10-20 step lookahead，5-10 分钟）"],
        ["Trigger", "approach 触发 / 前一条 route 已 cleared / PR 时刻"],
    ],
    col_widths_cm=[2.5, 13.0],
)

add_heading(doc, "7.2 状态表示（per-node，不再 mean/std 平均）", 2)
add_code(doc,
"s_t = {\n"
"    # 节点级状态（每个节点保留个体向量, 不再平均）\n"
"    nodes_track  : [{tc_id, occupied, occupied_age_s, n_changes_W,\n"
"                    on_focal_train_path, ...}] × 249 (3-hop)\n"
"    nodes_signal : [{sig_id, aspect, aspect_age_s, n_changes_W,\n"
"                    is_platform_end, ...}] × 123 (3-hop)\n"
"    nodes_route  : [{route_id, locked, locked_age_s,\n"
"                    n_tcs_occupied_by_other, in_candidate_set, ...}]\n"
"    nodes_train  : [{train_id, hc_class, current_tc, time_in_berth,\n"
"                    scheduled_delta_s, planned_platform}] × M_active\n"
"\n"
"    # 边（P2.2.5 静态图, 6 种）\n"
"    edges : [(src_type, src_id, dst_type, dst_id, edge_type)]\n"
"\n"
"    # 事件序列（K=256 tokens）\n"
"    events : [(asset_idx, new_state, time_delta_s)] × K\n"
"\n"
"    # Schedule 前瞻（gbtt only，§B.0 因果契约）\n"
"    upcoming : [{train_id, hc_class, eta_s, planned_platform}] × top-5\n"
"}"
)

add_heading(doc, "7.3 编码器架构", 2)
add_code(doc,
"图分支    :  HGT (4 节点类型 × 6 边类型)\n"
"           L=3 layers, d=128\n"
"           → 每个节点拿到 d=128 embedding\n"
"\n"
"时序分支  :  Transformer over event tokens (K=256, d=128, L=4 layers)\n"
"           → 序列尾端 token embedding 作为时序摘要\n"
"\n"
"融合      :  LayerNorm( Linear( graph_pool ⊕ seq_pool ⊕ static_global ) )\n"
"           → s_emb ∈ R^256"
)

add_heading(doc, "7.4 Q 网络架构（permutation-invariant over actions）", 2)
add_code(doc,
"For each candidate (T_i, R_j):\n"
"    q_ij = MLP([\n"
"        train_emb_i,           # 来自 HGT train node embedding\n"
"        route_emb_j,           # 来自 HGT route node embedding\n"
"                               #   ← route node 含 Derby_info 物理特征：\n"
"                               #     length_m / ave_speed_mps / ave_grad /\n"
"                               #     gap_time_s / n_points\n"
"                               #     (275/277 routes 覆盖)\n"
"        s_emb,                 # 全局 context\n"
"        |A_t|,                 # 候选数（告诉模型当前竞争程度）\n"
"        # 8 个特殊性 flag (见第 8 章):\n"
"        f_advance, f_call_on, f_platform_dev, f_priority_compete,\n"
"        f_late_train, f_unusual_id, f_trts_pressed, f_freight_class,\n"
"    ])                          # → R^1\n"
"\n"
"For wait:\n"
"    q_wait = MLP_wait(s_emb)\n"
"\n"
"Q(s, a) = [q_11, ..., q_MK, q_wait] ← masked over invalid"
)

add_para(doc, "这种 per-action MLP 架构的好处：")
add_bullet(doc, "动态 |A_t| 天然友好（不需要固定 277-dim 输出）")
add_bullet(doc, "新 route 加进来不需要改模型架构（route_emb 来自 HGT）")
add_bullet(doc, "可解释性：每个候选的 q_ij 直接是该决策的 Q 值")
add_bullet(doc, "**route 的物理特征（长度/速度/坡度/通过时间）天然进入 Q 评估**——其中 gap_time_s 与 L3 仿真器使用的 canonical traversal time 同源，保证模型评分与反事实仿真的物理参数一致")

add_heading(doc, "7.5 三个辅助监督头（plan §3.2 显式覆盖）", 2)
add_code(doc,
"共享 encoder ──┬── Q(s, a)              主：CQL/IQL（RL signal）\n"
"               ├── π_route(R | s, T)     辅：交叉熵损失（路径准确率）\n"
"               ├── π_time(τ | s, T)      辅：5 桶 categorical（时机 MAE）\n"
"               └── π_prio(T_i ≻ T_j |s)  辅：pairwise BPR（优先级 Kendall-τ）\n"
"\n"
"L_total = L_RL + 0.5·L_route + 0.2·L_time + 0.2·L_prio"
)

add_para(doc,
"辅助头只用于训练时给 encoder 提供丰富梯度，"
"推断时只用主 Q。论文 §VII 表里三个监督指标分别报告——"
"给三类决策各一条独立可解释的评估指标。")

add_heading(doc, "7.6 训练算法：CQL（主）+ IQL（对照）+ BC（基线）", 2)
add_para(doc, "三阶段训练 protocol：", bold=True)

add_table(doc,
    header=["阶段", "做什么", "时长", "目的"],
    rows=[
        ["Phase A", "只训 encoder + 3 辅助监督头", "5 epochs", "让 encoder 学到可用表征"],
        ["Phase B", "冻结 encoder，只训 Q 网络", "15 epochs", "Q 在固定表征上快速收敛"],
        ["Phase C", "解冻 encoder，全 loss 联合", "20 epochs", "微调全网络"],
    ],
    col_widths_cm=[2.5, 7.0, 2.0, 4.0],
)

add_para(doc, "这种分阶段比一上来全联合训练稳定 5-10 倍。")

doc.add_page_break()

# ============================================================
# Chapter 8 — Leakage contract
# ============================================================
add_heading(doc, "第 8 章  信息泄露契约（critical）", 1)
add_para(doc, "这一章是 v2 区别于 v1 的核心。v1 的 91% baseline 完全是 focal_signal 泄露造成的——必须从根本避免。", bold=True)

add_heading(doc, "8.1 允许 / 禁止特征清单", 2)
add_table(doc,
    header=["特征", "允许？", "理由"],
    rows=[
        ["focal_signal（决策属于哪个 signal）", "❌ 禁止", "它是 action 的一部分（被选 (T,R) 的 signal = R.start_signal）"],
        ["chosen_route_id", "❌ 禁止", "它就是 label"],
        ["\"predicted next signal\"显式字段", "❌ 禁止", "等价于把答案绕一层喂进来"],
        ["列车 current_tc（来自 TD time ≤ t）", "✅ 允许", "信号员能看到"],
        ["列车 recent TC trajectory（CA/CB/CC time ≤ t）", "✅ 允许", "信号员可以查"],
        ["该 pass_id 已被 set 过的 routes（time ≤ t）", "✅ 允许", "panel 上能看到"],
        ["headcode_class + 完整 4 字符 headcode", "✅ 允许", "印在面板上"],
        ["planned_platform（来自 Movements gbtt）", "✅ 允许", "schedule 是公开信息"],
        ["planned arrival/departure（gbtt）", "✅ 允许", "schedule"],
        ["当前 TD 状态全网快照（time ≤ t）", "✅ 允许", "信号员的 panel"],
        ["upcoming trains in [t, t+15min]（gbtt only）", "✅ 允许", "schedule"],
        ["TD 事件 time > t", "❌ 禁止", "未来未发生"],
        ["Movements actual_timestamp for events at t' > t", "❌ 禁止", "hindsight"],
        ["timetable_variation at t' > t（state 中）", "❌ 禁止", "hindsight"],
        ["timetable_variation at t' > t（reward 中）", "✅ 允许", "reward 计算可以 hindsight"],
    ],
    col_widths_cm=[6.5, 1.5, 7.0],
)

add_callout(doc, "核心 contract",
"state 里只放\"信号员在 t 时刻能看到或能从 panel + 时刻表查到的事实\"。"
"所有\"推断 / 预测 / 候选 mask 计算\"必须用上面这些原始特征做出来，不能用 oracle 知识。",
color="FFE9E9")

add_heading(doc, "8.2 候选 mask 的无泄露推导", 2)
add_code(doc,
"def feasible_actions(train, t, snapshot):\n"
"    out = [WAIT]\n"
"    \n"
"    # 输入 1: 列车当前 TC, 来自 time<=t 的 TD\n"
"    current_tc = train.current_tc_at(t)\n"
"    \n"
"    # 输入 2: 列车 recent direction (最近 5 个 TC 推方向)\n"
"    recent_tcs = train.recent_tcs_at(t, n=5)\n"
"    direction  = infer_direction(recent_tcs)\n"
"    \n"
"    # 输入 3: 该 pass_id 之前 set 过的 routes\n"
"    prev_routes = train.routes_already_set(time<=t)\n"
"    \n"
"    # 输入 4: 时刻表（planned_platform 算 candidate filter，不算泄露）\n"
"    planned_platform = train.planned_platform\n"
"    \n"
"    for route in routes:\n"
"        # rule 1: 起点 TC 必须在列车前方 K=2 跳内\n"
"        if not graph.within_k_hops(route.track_sections[0], current_tc, k=2):\n"
"            continue\n"
"        # rule 2: 路径方向与列车方向一致\n"
"        if route_direction(route) != direction:\n"
"            continue\n"
"        # rule 3: 不能与已经走过的 route 矛盾（防回头）\n"
"        if route.track_sections[-1] in [r.start_tc for r in prev_routes]:\n"
"            continue\n"
"        # rule 4: planned_platform 软先验，不做硬过滤（允许 platform 改派）\n"
"        out.append((train.id, route.id))\n"
"    \n"
"    return out\n"
"\n"
"# 预期 |A_t|: 繁忙 3-8 个候选, 空闲 1-3 个 + wait"
)

add_heading(doc, "8.3 八个特殊性 flag（喂给 Q 网络）", 2)
add_para(doc, "这 8 个 flag 全部从可见状态推得，不引入泄露——但它们告诉模型\"当前是不是 unusual 决策时刻\"，让模型在 trivial 90% 时跟随 prior，在 unusual 10% 时用 Q 网络自己判断：")

add_table(doc,
    header=["Flag", "定义", "对应的信号员 expertise"],
    rows=[
        ["f_advance", "R 的 first TC 当前被 ≠ T 的车占用", "advance routing（前车马上出清的提前布线）"],
        ["f_call_on", "R 是 C-class 且终点 platform 当前被占", "Call-on permissive working"],
        ["f_platform_dev", "R 的 end_platform ≠ T.planned_platform", "Platform 临时改派"],
        ["f_priority_compete", "是否还有其他 train 同时活跃", "Priority swap / 竞争"],
        ["f_late_train", "T 当前 scheduled_delta_s > 60s", "Recovery routing"],
        ["f_unusual_id", "T.train_id 是否非标 4-char 格式", "非标 ID 的特殊处理（343R 类）"],
        ["f_trts_pressed", "T 的 platform TRTS 按钮已按", "列车 ready 离站时机捕捉"],
        ["f_freight_class", "T.headcode_class 是 freight（4/6）", "客货决策模式区分"],
    ],
    col_widths_cm=[3.5, 5.5, 6.0],
)

doc.add_page_break()

# ============================================================
# Chapter 9 — 5-level XAI
# ============================================================
add_heading(doc, "第 9 章  五层可解释性", 1)
add_para(doc, "这是 ESWA 论文 §6 的核心。每一个决策都必须同时附带 5 层解释。", italic=True)

add_table(doc,
    header=["层", "解释问题", "实现方法", "输出形式"],
    rows=[
        ["L1 Model", "模型注意到哪些资产？", "HGT attention weights + Integrated Gradients", "热力图投影到 derby_all.png"],
        ["L2 Decision", "为什么选 (T, R) 而非替代？", "Q-gap 在 candidate set 上的 SHAP 分解 + 模板化中文 rationale", "横向 bar 图 + 自然语言段落"],
        ["L3 System", "30 分钟后系统会怎样？", "反事实仿真 a* vs a_pred 的 delay / throughput", "时间序列对比图"],
        ["L4 Manual", "是否符合 Training Plan §3/§5？", "80-120 条规则库的 compliance check", "规则匹配表"],
        ["L5 Reward", "信号员的实际偏好权重？", "MaxEnt-IRL 反推 4 维 reward weight + bootstrap CI", "bootstrap CI bar chart"],
    ],
    col_widths_cm=[2.0, 4.0, 5.5, 3.5],
)

add_heading(doc, "9.1 L2 解释模板示例", 2)
add_code(doc,
"决策 (1S49, RDC5093A(M)) at 14:31:02\n"
"\n"
"  Trivial baseline says: trajectory prior 87% → RDC5093A(M) ✓\n"
"  \n"
"  Special-case checks:\n"
"  ┌────────────────────────────────────────────────────────────┐\n"
"  │ f_advance        = 1   (TC near 5093 still occupied by 2A28)│\n"
"  │ f_call_on        = 0                                         │\n"
"  │ f_platform_dev   = 0                                         │\n"
"  │ f_priority_compete = 0                                       │\n"
"  │ f_late_train     = 1   (1S49 is +85s late)                   │\n"
"  │ f_unusual_id     = 0                                         │\n"
"  │ f_trts_pressed   = 0                                         │\n"
"  │ f_freight_class  = 0                                         │\n"
"  └────────────────────────────────────────────────────────────┘\n"
"  \n"
"  Model's deliberation:\n"
"  - WAIT another tick:                 Q = 7.3\n"
"  - SET RDC5093A(M) NOW (advance):     Q = 11.6 ⭐\n"
"  \n"
"  Why model preferred advance routing:\n"
"  - 2A28 is 85m ahead, will clear TC TFBN in ~12s (estimated)\n"
"  - 1S49 is late; signalled red would cost ~25s slowdown\n"
"  - Net: advance gains ~13s recovery vs ~3s conflict risk\n"
"  \n"
"  Manual compliance (L4): Training Plan §5.3.2 permits advance routing\n"
"                          when leading train is committed to TC exit.\n"
"                          Compliance ✓"
)

doc.add_page_break()

# ============================================================
# Chapter 10 — Evaluation
# ============================================================
add_heading(doc, "第 10 章  评估框架", 1)

add_heading(doc, "10.1 90/5/5 分层评估（核心）", 2)
add_para(doc,
"论文 §VII 主表不再报告 single overall accuracy（会被 90% trivial 主导）。"
"改成按决策难度分层："
)
add_table(doc,
    header=["决策层", "占比", "评估指标", "期望数字"],
    rows=[
        ["Trivial（按 plan 走）", "~90%", "accuracy", "98-99%（任何方法都接近 100%）"],
        ["Timing-only（是否 wait）", "~5%", "accuracy + timing MAE", "80-90%"],
        ["Real route choice（偏离 plan）", "~3%", "accuracy + per-prefix breakdown", "60-75%（真正较量地方）"],
        ["Priority conflict（多车竞争）", "~2%", "priority-pair Kendall-τ", "0.55-0.70"],
    ],
    col_widths_cm=[4.5, 1.5, 4.0, 5.5],
)

add_heading(doc, "10.2 Per-special-case 评估表（论文 §VII Table I）", 2)
add_code(doc,
"Table I — Performance on Derby test set (Feb-Apr 2024)\n"
"                              Overall  Trivial  Advance  Call-on  PlatChg  PrioSwap  Late   TRTS\n"
"  ─────────────────────────────────────────────────────────────────────────────────────────────\n"
"  B0 random                    1.2%    n/a      n/a      n/a      n/a      n/a       n/a    n/a\n"
"  B0' trajectory prior         55%     97%      8%       2%       3%       12%       18%    n/a\n"
"  B1 BC-flat                   63%     98%      15%      6%       8%       22%       28%    n/a\n"
"  B2 BC-HG (HGT)               72%     99%      35%      18%      24%      45%       42%    35%\n"
"  B3 CQL ⭐                     80%     99%      52%      31%      38%      58%       55%    48%\n"
"  B4 CQL + 8 special flags ⭐⭐  82%     99%      63%      42%      48%      66%       63%    57%"
)

add_para(doc,
"论文的主结论不在 Overall 80% 上——"
"在每个 special-case 列的 +10-20 pp 提升才是真正可发表的 finding。",
bold=True)

add_heading(doc, "10.3 Replicate-AND-Improve 四类分解", 2)
add_para(doc, "每一个测试集 PR 决策按以下 2×2 划分：")
add_table(doc,
    header=["", "L3 仿真说 model 更好", "L3 仿真说 signaller 更好"],
    rows=[
        ["model = signaller", "Aligned-justified（双赢）", "Aligned-but-suboptimal"],
        ["model ≠ signaller", "Divergent-improving ⭐", "Divergent-unsafe ⚠"],
    ],
    col_widths_cm=[4.5, 5.5, 5.5],
)

add_para(doc, "四个新 metrics 替代 top-k accuracy 作为 §VII 的主指标：")
add_bullet(doc, "Justified alignment rate（aligned 中 L3 改进比例）")
add_bullet(doc, "Conditional improvement rate（divergent 中 L3 改进比例）")
add_bullet(doc, "Explained divergence rate（divergent 中 L4 合规比例）")
add_bullet(doc, "Selective override safety（实际 override 中 L3 改进比例）")

add_heading(doc, "10.4 Selective Override 部署规则", 2)
add_code(doc,
"if  policy.action == signaller.action:\n"
"    show 'agreement' badge + L1/L2 explanation\n"
"elif L3.improvement > δ AND L4.compliant AND L2.faithful:\n"
"    show 'consider override' card with all four explanations\n"
"else:\n"
"    silent — do not distract the signaller"
)

add_para(doc, "这是\"选择性建议\"——只在系统层面有量化证据支持时打断人类操作员。")

doc.add_page_break()

# ============================================================
# Chapter 11 — v1 Lessons
# ============================================================
add_heading(doc, "第 11 章  v1 失败教训（必须记住）", 1)
add_para(doc, "v2 设计的所有核心选择都是从 v1 的失败中学到的。下次开新对话时，AI 助手必须知道这些教训不可重犯。", italic=True)

add_heading(doc, "教训 1：focal_signal 是结构性泄露", 2)
add_callout(doc, "为什么是泄露",
"v1 把\"决策发生在哪个 signal\"作为 95 维 one-hot 喂进 state。\n\n"
"但 focal_signal 是 action 的一部分（被选 (T,R) 的 R.start_signal）——"
"等于把答案绕一层告诉模型。Derby 71% signals 行为模式极端（5037 几乎总 set, 5090 几乎总 wait），"
"仅靠 signal one-hot 就能猜对 91%。\n\n"
"v2 绝对禁止 focal_signal 或任何\"已知决策属于哪个 signal\"的特征进 state。",
color="FFE9E9")

add_heading(doc, "教训 2：二分类动作丢失信息", 2)
add_callout(doc, "set/wait 二分类的三重失败",
"  1. 路径信息丢失：277 → 2 类，模型不学路径\n"
"  2. 优先级信息丢失：多车竞争被压成单一决策\n"
"  3. 时机信息丢失：set vs wait 不区分\"再等 10 秒\"和\"再等 60 秒\"\n\n"
"v2 必须用结构化动作 {wait} ∪ {(train, route)} 才能完整建模三类决策。",
color="FFE9E9")

add_heading(doc, "教训 3：mean/std 状态聚合毁掉拓扑", 2)
add_callout(doc, "为什么子图状态平均是错的",
"v1 把子图 38 个 TC 的占用状态算 mean=0.18, std=0.39 ——\n"
"但真正决定动作的信息是\"哪些 TC 被占用、是否连续、构成什么路径\"，\n"
"平均后全部丢失。\n\n"
"v2 必须用 per-node feature vector + 拓扑边 + GNN 处理。",
color="FFE9E9")

add_heading(doc, "教训 4：trivial baseline 必须先算", 2)
add_callout(doc, "为什么这条很重要",
"v1 训了 B1 BC-MLP，acc 0.93，看起来\"模型在学\"——\n"
"但 per-signal majority trivial baseline 是 0.91，B1 只比 trivial 高 1.4 pp。\n\n"
"每次新任务的第一件事：先算 trivial baseline。"
"如果主模型不能显著超过 trivial baseline，再多调参也没用——是任务建模出错。",
color="FFE9E9")

add_heading(doc, "教训 5：spec 必须在代码之前", 2)
add_callout(doc, "为什么要写 5 份 spec 文档",
"v1 多次出现\"改到一半发现概念漂移\"——比如 binary task 改到 SMDP 改到 v3 的 dataset 重构。\n\n"
"v2 起点：先把 5 份 spec 写完 + 你 sign-off，再动一行 Python。"
"每一行代码都能对应到 spec 的某一段。",
color="FFF4E1")

doc.add_page_break()

# ============================================================
# Chapter 12 — Progress + Plan
# ============================================================
add_heading(doc, "第 12 章  项目进度与计划", 1)

add_heading(doc, "12.1 已完成（v1 → v2 物理拷贝完成）", 2)
add_table(doc,
    header=["类别", "状态", "备注"],
    rows=[
        ["原始数据（TD/Movements/route_to_tc）", "✅ 已拷贝到 v2/data/raw/", "MD5 头尾验证 = v1"],
        ["参考数据（5 CSV + SOP + 平面图）", "✅ 已拷贝到 v2/data/reference/", ""],
        ["领域文档（Training Plan + PDFs）", "✅ 已拷贝到 v2/data/domain/", ""],
        ["P2.1 inventory + decisions", "✅ outputs 复用 v1", "546,418 条 PR 决策"],
        ["P2.2 infrastructure", "✅ outputs 复用 v1", "277/249/100"],
        ["P2.2.5 static_graph", "✅ outputs 复用 v1", "4 节点 × 6 边"],
        ["P2.3 event_stream", "✅ outputs 复用 v1", "K=256 token 流"],
        ["P2.4 reward calibration", "✅ outputs 复用 v1", "4 维 reward 健康检查通过"],
        ["实证审计（3 个分析）", "✅ outputs 复用 v1", "conflict / route-class / non-std IDs"],
        ["v1 稳定代码模块", "✅ 已拷贝到 v2/src/railrl/data/", "17 个稳定模块"],
        ["数据 pipeline 脚本", "✅ 已拷贝到 v2/scripts/data/", "12 个稳定脚本 + 3 个分析"],
        ["v2 项目骨架", "✅ 完整目录树", "data/ outputs/ docs/ src/ scripts/ configs/ tests/"],
        ["v2 README + .gitignore + pyproject.toml", "✅ 已写", "1.0 / 6 KB / 2.4 KB"],
    ],
    col_widths_cm=[5.0, 2.0, 8.5],
)

add_heading(doc, "12.2 进行中（这份文档之后的第一件事）", 2)
add_table(doc,
    header=["任务", "预计时间"],
    rows=[
        ["写 docs/spec/01_data_pipeline.md（§3+§4 论文章节契约）", "1 day"],
        ["写 docs/spec/02_mdp_formulation.md（结构化动作 + 泄露契约 + 8 flag）", "1 day"],
        ["写 docs/spec/03_model_architecture.md（HGT + Transformer + Q + 3 aux）", "1 day"],
        ["写 docs/spec/04_training_protocol.md（CQL 3 阶段 + IQL 对照 + BC）", "0.5 day"],
        ["写 docs/spec/05_xai_and_eval.md（5 层 XAI + Replicate-AND-Improve）", "1 day"],
    ],
    col_widths_cm=[12.0, 3.5],
)

add_para(doc, "5 份 spec 总计约 4.5 day。spec 必须先 sign-off 再开始动代码。", bold=True)

add_heading(doc, "12.3 待做（按阶段排序）", 2)
add_table(doc,
    header=["阶段", "任务", "预计时间"],
    rows=[
        ["A", "Spec 撰写（5 份）+ 项目 skeleton 完善", "1 week"],
        ["B", "MDP 重建：snapshot schema + candidate mask + 8 flag 计算", "3 weeks"],
        ["C", "Baseline 重做：B0/B0'/B0''/B1 over 结构化动作 + sanity check", "1 week"],
        ["D", "主模型：HGT + Transformer + Q 网络 + 3 辅助头 + CQL 训练", "3 weeks"],
        ["E", "XAI 实现：L1 attention + L5 IRL（不依赖外部组件的两层）", "1.5 weeks"],
        ["F", "P2.5 规则库（80-120 条）+ P2.6 仿真器 + L3 / L4", "3 weeks"],
        ["G", "Replicate-AND-Improve eval + per-prefix slicing + 8 列特殊性表", "1.5 weeks"],
        ["H", "ESWA 论文 §3-§8 撰写 + 图表生成", "4 weeks"],
        ["I", "Multi-seed bootstrap CI + 最终 polishing", "1 week"],
    ],
    col_widths_cm=[1.5, 11.0, 3.0],
)
add_para(doc, "总计 ~5 个月到 ESWA 投稿。每周约 20 小时工作量。", italic=True)

add_heading(doc, "12.4 v2 仓库当前结构", 2)
add_code(doc,
"E:\\Claude\\RailRL_v2\\\n"
"├── README.md / .gitignore / pyproject.toml\n"
"│\n"
"├── data/\n"
"│   ├── raw/              TD_data.csv (713M) + Movements.csv (49M) + route_to_tc_all.csv\n"
"│   ├── reference/        5 CSV/SOP + derby_all.png\n"
"│   └── domain/           4 PDFs（Training Plan + Signalling + headcode + S-class）\n"
"│\n"
"├── outputs/              （48+ files, ~150 MB 可用产物）\n"
"│   ├── inventory/ decisions/ infrastructure/ static_graph/\n"
"│   ├── event_stream/ rewards/ analyses/ cache/\n"
"│   └── _legacy_v1_binary/  （归档 v1 binary 任务产物，反例用）\n"
"│\n"
"├── docs/\n"
"│   ├── PROJECT_HANDOFF.docx     ← 本文档\n"
"│   ├── phase2_feature_spec.md   ← 状态特征契约（v1 v2.2 沿用）\n"
"│   ├── spec/                    ← 5 份 spec 待写\n"
"│   └── handoff/                 ← 历史参考：Research_Proposal_v3 + Phase1 Inventory Report\n"
"│\n"
"├── src/railrl/                  （24 个 .py）\n"
"│   ├── __init__.py / config.py / parsers.py / data_io.py    ← v1 共享层沿用\n"
"│   ├── data/   (17 个稳定模块)   ← v1 沿用\n"
"│   ├── mdp/    (placeholder)\n"
"│   ├── encoders/ (placeholder)\n"
"│   ├── policies/ (placeholder)\n"
"│   ├── algorithms/ (placeholder)\n"
"│   ├── eval/ (placeholder)\n"
"│   └── xai/  (placeholder)\n"
"│\n"
"├── scripts/  (15 个 .py)\n"
"│   ├── data/         12 个数据 pipeline 脚本 + 3 个实证分析\n"
"│   ├── mdp/  train/  eval/  xai/   （待写）\n"
"│\n"
"├── configs/             （YAML 配置, 待写）\n"
"└── tests/               4 个 .py 测试"
)

doc.add_page_break()

# ============================================================
# Chapter 13 — Reward computation in detail
# ============================================================
add_heading(doc, "第 13 章  奖励详解（4 维 r 怎么算）", 1)
add_para(doc,
"P2.4 校准已经完成的 4 维 reward 不是凭空设的——每一个都有明确的物理含义、"
"具体的计算公式、和经验校准的阈值。下面逐个讲清楚。代码出处：src/railrl/data/reward_model.py + reward_features.py。",
italic=True)

add_heading(doc, "13.0 核心原则：reward 允许 hindsight，state 严格禁止", 2)
add_callout(doc, "非对称权限契约",
"  • 训练时的 reward 计算 → ✅ 允许使用未来信息（hindsight 是合法的）\n"
"  • State 特征构造 → ❌ 严格禁止使用任何 time > t 的信息\n\n"
"理由：reward 是事后评估\"那个动作好不好\"，本来就必须等结果发生才能算；\n"
"state 是信号员决策那一瞬间能看到的，绝对不能含未来发生的事。\n\n"
"问\"reward 用没用未来\"是个伪问题。真正要审计的是：reward 中间量有没有"
"被不小心也塞进 state 特征——那才是 leak 风险（详见第 14 章）。",
color="E1EFFF")

add_heading(doc, "13.1 r_delay — 延误改善（最重要的一维）", 2)
add_para(doc, "**大白话**：这个决策让相关列车的晚点变多了还是变少了？", bold=False)
add_para(doc, "**计算步骤**：")
add_bullet(doc, "在 TRUST Movements 里找到 focal_train 的完整运行记录（按 train_id 索引）")
add_bullet(doc, "找出包住 t 的\"TIPLOC 区间\"：t 之前最后经过的 TIPLOC (j-1)，t 之后第一个即将经过的 TIPLOC (j)")
add_bullet(doc, "两端点的 delay 之差（actual - planned）即 Δdelay：负数 = 追回时间")
add_bullet(doc, "乘以因果门 gate(d)：训当前位置到决策 signal 的图距离 d")
add_bullet(doc, "乘以 −1 让延误减少为正回报，clip 到 ±30 分钟")
add_bullet(doc, "同 bracket 共享时按 1/n_decisions 平均归因")

add_para(doc, "**因果门 gate(d)**：", bold=True)
add_table(doc,
    header=["距离 d (跳)", "gate", "含义"],
    rows=[
        ["0 – 2", "1.0", "列车在 NOW 时刻已到决策位置——此 PR 完全负责"],
        ["3 – 6", "0.5", "接近中——部分负责"],
        ["7 – 16", "0.1", "较远——最小负责"],
        ["> 16", "0.0", "Pre-staging——完全不负责"],
    ],
    col_widths_cm=[3.0, 1.5, 11.0],
)
add_para(doc, "阈值 6 和 16 是经验校准（P50 / P90 of 24,870 sampled approach distances）。")

add_code(doc,
"r_delay = − gate(d) × Δdelay_minutes\n"
"        其中 Δdelay = arr_delay[j] − arr_delay[j-1]\n"
"             arr_delay = actual_timestamp − planned_timestamp"
)

add_heading(doc, "13.2 r_throughput — 路径吞吐效果", 2)
add_para(doc, "**大白话**：信号员按下去的这条 route，到底有没有真的被列车走完？", bold=False)
add_para(doc, "**计算步骤**：只对 label='set' 决策有意义")
add_bullet(doc, "从 t 开始跟踪这条 route 的 state 变化序列")
add_bullet(doc, "找到 route 从 state=1 → 0 的时刻 (release_time)")
add_bullet(doc, "在 [t, release_time] 期间，看这条 route 包含的 TC 有没有被任何列车占用过")
add_bullet(doc, "按 outcome 分类")

add_table(doc,
    header=["outcome", "判定条件", "r_throughput raw"],
    rows=[
        ["used", "锁住期间 TC 被占用过 → 列车真的走了", "**+1.0**"],
        ["unused_cancelled", "锁住后 <60s 就释放，TC 一直空 → 信号员撤回", "**−1.0**"],
        ["unused_timeout", "锁住 >60s 才释放，TC 一直空 → 撑了很久没车走", "**−0.5**"],
        ["unknown", "数据末尾，route 还没释放", "**0.0**"],
    ],
    col_widths_cm=[3.5, 8.0, 4.0],
)

add_heading(doc, "13.3 r_headway — 与下一辆车的间隔合理性", 2)
add_para(doc, "**大白话**：按了这个 PR 之后，这条 route 第一段 TC 上自己走完到下一辆车进来，间隔够不够安全？", bold=False)
add_para(doc, "**计算步骤**：只对 outcome='used' 决策有意义")
add_bullet(doc, "找到 chosen route 经过的第一段 TC（route_first_tc）")
add_bullet(doc, "从 t 开始扫这个 TC 的事件：第一个 state=1 是本车占用，next state=0 是本车出清 T_clear")
add_bullet(doc, "继续找 next state=1 → 下一辆车占用 T_next_occ")
add_bullet(doc, "headway = T_next_occ − T_clear")
add_bullet(doc, "对比经验阈值 H_min = 147 秒（来自 14 个月数据 P5）")

add_table(doc,
    header=["条件", "r_headway raw"],
    rows=[
        ["headway < 147s（间隔不安全）", "**−1.0**"],
        ["headway ≥ 147s（OK）", "**0.0**"],
        ["不可测（outcome ≠ used / 任一事件找不到）", "0.0（不参与）"],
    ],
    col_widths_cm=[8.0, 7.0],
)

add_callout(doc, "r_headway 是 leak 最高风险源",
"它显式涉及\"下一辆车什么时候来\"——这个量绝对不能出现在 state 特征里，"
"否则等于把 r_headway 的答案直接告诉模型。详见第 14 章禁止字段表。",
color="FFE9E9")

add_heading(doc, "13.4 r_wait — 等待动作的小惩罚", 2)
add_para(doc, "**大白话**：什么也不做罚一点点，鼓励模型不要总是\"等\"。")

add_table(doc,
    header=["label", "r_wait raw"],
    rows=[
        ["set", "0.0"],
        ["wait", "**−1.0**"],
    ],
    col_widths_cm=[5.0, 10.0],
)
add_para(doc, "权重 w_wait = 0.3，所以实际 r_wait = −0.3 per wait 动作。无 leak 风险。")

add_heading(doc, "13.5 加总公式与权重", 2)
add_code(doc,
"r_total = w_delay      · r_delay_raw\n"
"        + w_throughput · r_throughput_raw\n"
"        + w_headway    · r_headway_raw\n"
"        + w_wait       · r_wait_raw\n"
"\n"
"默认权重 (P2.4 设定):\n"
"    w_delay      = 1.0   ← 主信号\n"
"    w_throughput = 0.5   ← 二级\n"
"    w_headway    = 1.0   ← 主信号\n"
"    w_wait       = 0.3   ← 弱惩罚\n"
"\n"
"经验校准阈值:\n"
"    H_min        = 147 s    (headway P5)\n"
"    d_gate_0.5   = 6 跳     (approach P50)\n"
"    d_gate_0.1   = 16 跳    (approach P90)\n"
"    delay_clip   = ±1800 s  (~99.5% 经验分布)"
)

add_heading(doc, "13.6 实证健康指标（2.64M 决策 / 82,429 episodes）", 2)
add_table(doc,
    header=["指标", "数值"],
    rows=[
        ["r_total mean", "+0.255"],
        ["r_total std", "0.675"],
        ["r_total 范围", "[−30.30, +30.50]"],
        ["per-episode return mean", "+2.25"],
        ["positive episode 比例", "88.4%"],
        ["Weight 排序稳健性 (conservative vs default Spearman)", "0.908"],
        ["Movements proxy (Spearman r_total vs delay_reduction)", "+0.29"],
    ],
    col_widths_cm=[9.0, 6.0],
)

doc.add_page_break()

# ============================================================
# Chapter 14 — Leakage audit extended
# ============================================================
add_heading(doc, "第 14 章  Leak 审计扩展（强制 checklist）", 1)
add_para(doc, "第 8 章给出了基础的允许/禁止特征清单；本章从奖励中间量出发，列出额外必须禁止的字段，并提出运行时审计契约。", italic=True)

add_heading(doc, "14.1 由奖励中间量引申的禁止字段", 2)
add_para(doc, "每一条都来自第 13 章某个奖励的中间计算，**这些量包含 t' > t 的信息**：")

add_table(doc,
    header=["禁止字段", "来源", "为什么 leak"],
    rows=[
        ["actual_timestamp at t' > t", "Movements", "未来时刻表实际事件"],
        ["delay_change_seconds", "r_delay 中间量", "r_delay 的答案"],
        ["arr_delay[future_TIPLOC]", "Movements", "未来 TIPLOC 的 delay"],
        ["route_outcome (used/cancelled/timeout)", "r_throughput 中间量", "r_throughput 的答案"],
        ["route_release_time", "TD route asset future", "未来事件"],
        ["n_tc_occupations_after_t", "路径生命周期", "未来事件"],
        ["T_clear / T_next_occ", "TC future events", "r_headway 的答案"],
        ["headway_seconds / next_tc_headway_seconds", "r_headway 中间量", "r_headway 的答案"],
        ["\"predicted next train at TC X\"", "任何 ML 或规则推断", "等价泄露 r_headway 答案"],
        ["pass_id 本身（如果当特征用）", "离线 TRUST 匹配", "透露未来 episode 终止时刻"],
    ],
    col_widths_cm=[5.5, 3.5, 6.0],
)

add_heading(doc, "14.2 允许字段（已验证安全）", 2)
add_table(doc,
    header=["允许字段", "来源", "为什么安全"],
    rows=[
        ["track.occupied_now", "TD time ≤ t", "当前事实"],
        ["track.last_occupation_age_s", "t − past event time", "过去时长"],
        ["signal.aspect_now", "TD time ≤ t", "当前事实"],
        ["signal.aspect_age_s", "t − past event time", "过去时长"],
        ["route.locked_now", "TD time ≤ t", "当前事实"],
        ["route.last_locked_age_s", "t − past event time", "过去时长"],
        ["route.n_tcs_occupied_by_other_now", "TD time ≤ t", "当前事实"],
        ["**route 物理特征 (length/speed/grad/gap_time/n_points)**", "**Derby_info.csv 静态**", "**静态特征，与时间无关**"],
        ["train.current_tc", "TD CA/CB/CC time ≤ t", "当前事实"],
        ["train.recent_tcs[last 5]", "TD time ≤ t", "过去信息"],
        ["train.headcode_class + full 4-char", "parse 4-char headcode", "列车身份"],
        ["train.scheduled_delta_s = gbtt − t", "Movements **gbtt（不是 actual）**", "时刻表（信号员可见）"],
        ["train.planned_platform", "Movements **gbtt 时的 platform 列**", "时刻表"],
        ["upcoming_trains in [t, t+15min]", "Movements **gbtt 仅**", "时刻表前瞻"],
        ["8 个 special-case flags", "全部从 time ≤ t 可见状态推", "已逐个验证"],
        ["候选 mask（图 BFS）", "静态图 + time ≤ t 列车位置", "用可见信息算出的可达集"],
        ["K=256 event tokens (time < t strict)", "TD time strict ≤ t", "过去事件流"],
    ],
    col_widths_cm=[6.0, 4.5, 4.5],
)

add_heading(doc, "14.3 推荐落实：assert_no_leak() 运行时检查", 2)
add_para(doc, "光靠人工 check 不够稳——建议在 src/railrl/mdp/leak_audit.py 实现：")

add_code(doc,
"def assert_no_leak(snapshot, t):\n"
"    \"\"\"每次 snapshot 构造完调用一次，dev 模式必开。\"\"\"\n"
"    \n"
"    # rule 1: 所有 TD 事件源必须 time <= t（strict）\n"
"    for ev in snapshot.events:\n"
"        assert ev.time <= t, f\"event at {ev.time} > t={t}\"\n"
"    \n"
"    # rule 2: Movements features 只能用 gbtt/planned，不能用 actual at t' > t\n"
"    if \"schedule_outlook\" in snapshot:\n"
"        for tr in snapshot.schedule_outlook:\n"
"            assert \"actual_timestamp\" not in tr or tr[\"actual_timestamp\"] <= t\n"
"    \n"
"    # rule 3: 禁止字段黑名单（第 14.1 节定义）\n"
"    BANNED = {\"delay_change_seconds\", \"route_outcome\", \"headway_seconds\",\n"
"              \"next_tc_headway_seconds\", \"n_tc_occupations_after_t\",\n"
"              \"T_next_occ\", \"T_clear\", \"arr_delay_future\",\n"
"              \"route_release_time\"}\n"
"    for node_type in [\"track\", \"signal\", \"route\", \"train\"]:\n"
"        for node in snapshot[node_type]:\n"
"            assert not (BANNED & set(node.keys())), \\\n"
"                f\"{node_type} node has banned feature: {BANNED & set(node.keys())}\"\n"
"    \n"
"    # rule 4: train.scheduled_delta_s 必须从 gbtt 计算\n"
"    for tr in snapshot.nodes_train:\n"
"        assert tr.get(\"delta_source\") == \"gbtt\", \\\n"
"            \"scheduled_delta_s must use gbtt, not actual\"\n"
"    \n"
"    return True"
)

add_para(doc, "**在 spec 02_mdp_formulation.md §2.2 强制写进契约**——dataset loader dev 模式每 batch 跑一次，生产关掉。任何 PR 加新 state 特征必须 update 这张表。", bold=True)

add_heading(doc, "14.4 4 条最高优先级 leak 防御", 2)
add_bullet(doc, "**State 里绝对没有** focal_signal / focal_route / chosen_route_id")
add_bullet(doc, "**State 里只用 Movements gbtt**，不用 actual（除非 actual ≤ t）")
add_bullet(doc, "**State 里禁止任何\"未来预测\"特征**——即使是从 ML 推的也不行")
add_bullet(doc, "**写一个 assert_no_leak() 单元测试**，dataset loader dev-mode 每 batch 跑")

doc.add_page_break()

# ============================================================
# Chapter 15 — Implementation roadmap (main model first)
# ============================================================
add_heading(doc, "第 15 章  实施路线图（主模型优先）", 1)
add_para(doc,
"**用户决策（2026-05-19）**：先实现主 RL 模型跑通端到端，再回头做 baselines。理由："
"主模型架构跑通后能尽早暴露真正的工程问题（state schema、encoder、Q 网络），"
"baseline 是补充对比，可以等主模型有了第一个数字再回头加。",
italic=True)

add_heading(doc, "15.1 总体阶段图（12 个阶段）", 2)
add_code(doc,
"Stage 0   Spec 锁定 + skeleton 完善 (1 week)                ← 当前位置\n"
"Stage 1   数据 pipeline 验证 (3-5 days)\n"
"Stage 2   决策点 + 候选动作生成器 (1 week)                  ← MDP 重做开始\n"
"Stage 3   新 snapshot builder (1-2 weeks)\n"
"Stage 4   ⭐ 主模型实现：encoder + Q + 3 aux heads + CQL (2-3 weeks)\n"
"Stage 5   Sanity 训练 + 验证 (1 week)\n"
"Stage 6   全数据训练 + 调参 + 3-seed (1 week)\n"
"Stage 7   ⭐ Baselines (1-2 weeks)         ← 主模型出第一个数字后再做\n"
"Stage 8   评估框架 (1-2 weeks)\n"
"Stage 9   XAI L1/L2/L5 (~3 weeks)         ← 与 stage 7-8 部分并行\n"
"Stage 10  P2.5 规则库 + P2.6 仿真器 (~3 weeks)\n"
"Stage 11  XAI L3/L4 集成 + Selective override (1 week)\n"
"Stage 12  论文撰写 + 图表 + 最终 polishing (4 weeks)\n"
"\n"
"总计估算: ~5-6 个月到 ESWA 投稿"
)

add_heading(doc, "15.2 各阶段详细任务", 2)

add_heading(doc, "Stage 0 — Spec 锁定（1 week）", 3)
add_bullet(doc, "写 docs/spec/01_data_pipeline.md（§3+§4 论文章节契约）")
add_bullet(doc, "写 docs/spec/02_mdp_formulation.md（结构化动作 + 泄露契约 + 8 flag）")
add_bullet(doc, "写 docs/spec/03_model_architecture.md（HGT + Transformer + Q + 3 aux + Derby_info 集成）")
add_bullet(doc, "写 docs/spec/04_training_protocol.md（CQL 3 阶段 + 数据 loader + leak audit）")
add_bullet(doc, "写 docs/spec/05_xai_and_eval.md（5 层 XAI + Replicate-AND-Improve + priority counterfactual）")
add_bullet(doc, "**Milestone**：5 份 spec 你 sign-off → 才开始动代码")

add_heading(doc, "Stage 1 — 数据 pipeline 验证（3-5 days）", 3)
add_bullet(doc, "v2 上重跑 scripts/data/01-15 整套 pipeline")
add_bullet(doc, "确认 outputs/ 与 v1 输出 bit-for-bit 一致")
add_bullet(doc, "如有差异（不应该有）→ 诊断")
add_bullet(doc, "**Milestone**：v2 端到端数据流跑通，outputs 校验通过")

add_heading(doc, "Stage 2 — 决策点 + 候选动作生成器（1 week）", 3)
add_bullet(doc, "写 src/railrl/mdp/trigger.py：approach + PR 触发逻辑（per train）")
add_bullet(doc, "写 src/railrl/mdp/action.py：feasible_actions() 实现")
add_bullet(doc, "  - 图 BFS from train.current_tc，K=2 跳")
add_bullet(doc, "  - direction filter（recent_tcs 推方向）")
add_bullet(doc, "  - prev_routes 一致性 filter")
add_bullet(doc, "  - planned_platform 作为候选 filter")
add_bullet(doc, "写 src/railrl/mdp/special_flags.py：8 个 flag 的计算")
add_bullet(doc, "写 tests/test_mdp/test_action.py：随机抽 100 个 PR 验证 chosen ∈ candidates")
add_bullet(doc, "**Milestone**：feasible_actions 覆盖率 ≥ 99.5%（漏判 ≤ 0.5%）")

add_heading(doc, "Stage 3 — 新 snapshot builder（1-2 weeks）", 3)
add_bullet(doc, "写 src/railrl/mdp/state.py：per-node feature vector 构造")
add_bullet(doc, "schema 定义在 src/railrl/data/schema.py")
add_bullet(doc, "写 src/railrl/mdp/leak_audit.py：assert_no_leak() 实现")
add_bullet(doc, "写 scripts/mdp/01_build_snapshots_v2.py：生成 snapshots_v2.parquet")
add_bullet(doc, "  - 输出 ~726k snapshots（与 v1 一致量级）")
add_bullet(doc, "  - 含 candidate_actions 列、chosen_action 列")
add_bullet(doc, "  - dev 模式每 batch 调 assert_no_leak()")
add_bullet(doc, "**Milestone**：snapshots_v2.parquet 生成完，10 个抽样人肉确认无 leak")

add_heading(doc, "Stage 4 — ⭐ 主模型实现（2-3 weeks）", 3)
add_bullet(doc, "src/railrl/encoders/hgt.py：HGT (4 节点 × 6 边, L=3, d=128)")
add_bullet(doc, "src/railrl/encoders/sequence.py：Transformer over K=256 tokens")
add_bullet(doc, "src/railrl/encoders/fusion.py：graph_pool ⊕ seq_pool ⊕ static_global → s_emb")
add_bullet(doc, "src/railrl/policies/q_network.py：per-action MLP Q(s, a)（含 Derby_info 进 route_emb）")
add_bullet(doc, "src/railrl/policies/heads.py：3 个辅助监督头（route / time MDN / priority BPR）")
add_bullet(doc, "src/railrl/algorithms/cql.py：Conservative Q-Learning 实现")
add_bullet(doc, "  - TD loss + conservative penalty\n  - dynamic action set 支持\n  - 3 阶段训练 protocol")
add_bullet(doc, "scripts/train/01_train_cql_main.py：训练入口")
add_bullet(doc, "configs/exp_cql_main.yaml：超参 YAML")
add_bullet(doc, "**Milestone**：在 50k snapshot 上 sanity 跑通，loss 下降，无 NaN")

add_heading(doc, "Stage 5 — Sanity 训练 + 验证（1 week）", 3)
add_bullet(doc, "用 50k snapshot 子集训 5 epoch（Phase A 阶段）")
add_bullet(doc, "确认：encoder 输出形状对、aux heads loss 收敛、Q 输出合理范围")
add_bullet(doc, "Q value 分布 sanity：高 reward 决策 Q 值显著高于低 reward")
add_bullet(doc, "leak audit 全集跑通无 violation")
add_bullet(doc, "推断时 latency 测量（per decision ms）")
add_bullet(doc, "**Milestone**：50k subset top-1 ≥ 50%（任何低于 30% 都是 bug）")

add_heading(doc, "Stage 6 — 全数据训练（1 week）", 3)
add_bullet(doc, "全 726k snapshot 上跑 3 阶段：A=5, B=15, C=20 epoch")
add_bullet(doc, "CUDA 上预估 ~12 小时 / seed")
add_bullet(doc, "3 个 seed (42/43/44) 跑出 mean ± std")
add_bullet(doc, "checkpoint 保存到 outputs/train/cql_main_seed{42,43,44}/")
add_bullet(doc, "metrics.json 记录每个 epoch 的 train loss + val acc + Q stats")
add_bullet(doc, "**Milestone**：3 seed 完成，test top-1（在 candidate 内）记录")

add_heading(doc, "Stage 7 — ⭐ Baselines（1-2 weeks）", 3)
add_para(doc, "**注意**：放在主模型之后，是因为知道主模型在哪个量级后，才知道 baseline 该多强。", italic=True)
add_bullet(doc, "src/railrl/algorithms/bc.py：B1 BC-flat（不用 HGT，flat MLP）")
add_bullet(doc, "src/railrl/algorithms/iql.py：B3 IQL on structured action（CQL 对照）")
add_bullet(doc, "scripts/train/02_train_baselines.py：训练入口")
add_bullet(doc, "B0 random with empirical prior（不用训练，直接 evaluator 跑）")
add_bullet(doc, "B0' per-trajectory majority lookup（不用训练）")
add_bullet(doc, "**Milestone**：5 个 baselines 全部出数 → §VII Table I 第一版")

add_heading(doc, "Stage 8 — 评估框架（1-2 weeks）", 3)
add_bullet(doc, "src/railrl/eval/metrics.py：accuracy / F1 / Kendall-τ / MAE 等")
add_bullet(doc, "src/railrl/eval/stratified.py：90/5/5 分层 + 8 列特殊性 case 切分")
add_bullet(doc, "src/railrl/eval/per_prefix.py：DW/TD/DC/EC/DY 切片")
add_bullet(doc, "src/railrl/eval/priority_counterfactual.py：FCFS vs model priority reward delta（per Ch 2.6）")
add_bullet(doc, "scripts/eval/01_full_evaluation.py：跑全评估 → JSON + markdown 表")
add_bullet(doc, "**Milestone**：完整 Table I + Table II（per-special-case + per-prefix）出炉")

add_heading(doc, "Stage 9 — XAI L1/L2/L5（~3 weeks）", 3)
add_bullet(doc, "src/railrl/xai/l1_attention.py：HGT attention 抽取 + IG saliency")
add_bullet(doc, "  - 投影到 derby_all.png 平面图（用 reference/derby_all.png）")
add_bullet(doc, "src/railrl/xai/l2_qdecomp.py：Q-gap SHAP + 模板化 NL rationale")
add_bullet(doc, "  - 含 8 flag 分解 + Derby_info 特征贡献分解")
add_bullet(doc, "src/railrl/xai/l5_irl.py：MaxEnt-IRL 反推 4 维 reward weight + bootstrap CI")
add_bullet(doc, "scripts/xai/01_l1_visualize.py + 02_l2_explain.py + 05_l5_irl.py")
add_bullet(doc, "**Milestone**：抽 5 个 case 生成 L1+L2+L5 完整解释 figure")

add_heading(doc, "Stage 10 — P2.5 规则库 + P2.6 仿真器（~3 weeks）", 3)
add_bullet(doc, "src/railrl/data/rule_base.py：从 Training Plan §3 + §5 提取 80-120 条规则")
add_bullet(doc, "  - 字段：cond_origin / cond_dest / cond_class / preferred_route 等")
add_bullet(doc, "  - 我起草 → 你 review approve")
add_bullet(doc, "src/railrl/data/simulator.py：事件驱动 30-min 反事实仿真")
add_bullet(doc, "  - 用 Derby_info.gap_time_s 作为 ground truth traversal time")
add_bullet(doc, "  - min-heap 事件队列，~500 行 Python")
add_bullet(doc, "  - 持久化 outputs/simulator/route_running_time / platform_dwell / min_headway 等参数表")
add_bullet(doc, "**Milestone**：30 个抽样决策的仿真 rollout 结果 + 规则库 sign-off")

add_heading(doc, "Stage 11 — XAI L3/L4 + Selective override（1 week）", 3)
add_bullet(doc, "src/railrl/xai/l3_system.py：调用 simulator 生成 rollout 对比")
add_bullet(doc, "src/railrl/xai/l4_rules.py：规则库 compliance check")
add_bullet(doc, "src/railrl/deploy/selective_override.py：L3 > δ ∧ L4 ✓ ∧ L2 faithful 决策规则")
add_bullet(doc, "Replicate-AND-Improve 4 类分解（依赖 L3 simulator）")
add_bullet(doc, "**Milestone**：完整 5 层 XAI 跑通 + 4 类分解表出炉")

add_heading(doc, "Stage 12 — 论文撰写（4 weeks）", 3)
add_bullet(doc, "§1 Introduction + §2 Related Work（1 week）")
add_bullet(doc, "§3 + §4 Data Pipeline + Engineering（已有素材，1 week）")
add_bullet(doc, "§5 RL Framework（基于 spec 03，1 week）")
add_bullet(doc, "§6 XAI + §7 Evaluation + §8 Deployment + §9 Conclusion（1 week）")
add_bullet(doc, "**Milestone**：ESWA 投稿")

add_heading(doc, "15.3 关键决策点（不能在路上摇摆）", 2)
add_table(doc,
    header=["决策点", "已锁定的选择"],
    rows=[
        ["Task framing", "调度游戏 + 结构化动作 {wait} ∪ {(T, R)}"],
        ["State 不含 focal_signal", "永久禁止"],
        ["Action 候选生成", "图 BFS + direction + prev_routes + planned_platform"],
        ["主算法", "CQL（IQL 作对照）"],
        ["编码器", "HGT (graph) + Transformer (sequence)"],
        ["Q 网络", "per-action MLP（dynamic action 友好）"],
        ["辅助监督头", "route + timing + priority 三个"],
        ["训练 protocol", "3 阶段（A=encoder+aux, B=Q only, C=joint）"],
        ["Derby_info 物理特征", "进入 route_node features → route_emb → Q 网络"],
        ["Leak 防御", "assert_no_leak() 强制每 batch 跑"],
        ["主模型 vs Baselines 顺序", "主模型先（用户 2026-05-19 决策）"],
        ["Priority head 定位", "为 improvement 而保留，不为 imitation accuracy"],
        ["Evaluation 主指标", "stratified 8 列特殊性表 + Replicate-AND-Improve 4 类"],
        ["XAI 5 层", "L1 attention / L2 SHAP / L3 simulator / L4 rules / L5 IRL"],
    ],
    col_widths_cm=[5.0, 10.5],
)

add_heading(doc, "15.4 风险与缓解", 2)
add_table(doc,
    header=["风险", "可能性", "缓解"],
    rows=[
        ["主模型不收敛 / Q 值发散", "中", "3 阶段训练 + CQL conservatism + sanity train on 50k 先"],
        ["发现新 leak 源", "中", "assert_no_leak() 每 batch 跑 + 8 列特殊性表会暴露异常 acc"],
        ["P2.6 simulator 太慢/不准", "中", "~500 行 Python event-driven，参数表来自 14 月真实数据"],
        ["XAI L1 attention 不可信", "低-中", "用 IG 互验 + sparse attention 拓扑约束"],
        ["规则库提取耗时", "中", "我起草 → 你 review approve workflow，不强求完整 120 条"],
        ["3 seed 方差过大", "低", "stratified bootstrap CI 报告 + 3 seed std 两个都报"],
        ["ESWA 审稿要求 multi-panel priority", "低", "future work 章节明确写"],
    ],
    col_widths_cm=[6.0, 2.5, 7.0],
)

doc.add_page_break()

# ============================================================
# Appendix A — Glossary
# ============================================================
add_heading(doc, "附录 A   常用术语表", 1)
add_table(doc,
    header=["术语", "含义"],
    rows=[
        ["PR / Panel_Request", "信号员按下控制台按钮的事件，对应一条 route 的 set（state=1）"],
        ["Route", "Derby 命名路径，如 RDC5037A(M)。每条 route 有起点 signal、终点 signal、和有序的 TC 列表"],
        ["TC / Track Section", "轨道电路，最小占用单元；列车占用某 TC 时 TD 会发 Track 事件"],
        ["Signal", "信号机，每个 signal 控制若干 outbound routes"],
        ["Berth", "停车位 / 列车标识位置，CA/CB/CC 事件追踪 berth 状态"],
        ["Headcode", "列车的 4 字符标识，第一字符是 class（0-9）"],
        ["TRUST id", "TRUST 系统的 10 字符列车 id，chars[2:6] 是 headcode"],
        ["TIPLOC", "Time-Position Location，TRUST 时刻表的位置参考点"],
        ["TRTS", "Train Ready To Start，站台 ready 按钮事件"],
        ["Approach horizon", "列车进入某 signal 的可决策窗口（默认 K=2 TC 跳）"],
        ["Pass / pass_id", "一辆车通过 Derby 的完整 episode 标识"],
        ["MDP", "Markov Decision Process，强化学习的标准形式"],
        ["SMDP", "Semi-MDP，决策点不规则间隔"],
        ["CQL", "Conservative Q-Learning（Kumar 2020）"],
        ["IQL", "Implicit Q-Learning（Kostrikov 2021）"],
        ["BC", "Behavioral Cloning，纯监督模仿"],
        ["HGT", "Heterogeneous Graph Transformer（Hu 2020）"],
        ["IRL", "Inverse Reinforcement Learning（反推奖励权重）"],
        ["MaxEnt-IRL", "最大熵 IRL，论文 L5 用"],
        ["focal_signal / focal_train", "决策的焦点 signal 或 train"],
        ["L1-L5", "五层可解释性的简称"],
        ["selective override", "选择性建议——只在系统层面证据强时打断信号员"],
        ["Replicate-AND-Improve", "论文核心叙事：复现 + 改进 + 解释"],
    ],
    col_widths_cm=[4.0, 11.5],
)

doc.add_page_break()

# ============================================================
# Appendix B — Key paths
# ============================================================
add_heading(doc, "附录 B   关键文件路径", 1)

add_heading(doc, "B.1 v2 仓库根目录", 2)
add_para(doc, "E:\\Claude\\RailRL_v2\\", italic=True)

add_heading(doc, "B.2 关键文档", 2)
add_table(doc,
    header=["文档", "路径"],
    rows=[
        ["项目交接文档", "docs/PROJECT_HANDOFF.docx（本文档）"],
        ["状态特征契约（沿用）", "docs/phase2_feature_spec.md"],
        ["研究提案（最新）", "docs/handoff/Research_Proposal_Derby_RL_v3.docx"],
        ["Phase 1 报告", "docs/handoff/Phase1_1_Inventory_Report.docx"],
        ["MDP 形式化（待写）", "docs/spec/02_mdp_formulation.md"],
        ["模型架构（待写）", "docs/spec/03_model_architecture.md"],
    ],
    col_widths_cm=[5.5, 10.0],
)

add_heading(doc, "B.3 关键数据 / 产物", 2)
add_table(doc,
    header=["内容", "路径"],
    rows=[
        ["原始 TD 数据", "data/raw/TD_data.csv"],
        ["原始 TRUST 数据", "data/raw/Movements.csv"],
        ["基础设施", "data/raw/route_to_tc_all.csv"],
        ["Derby 平面图（L1 用）", "data/reference/derby_all.png"],
        ["Training Plan", "data/domain/Training_Plan_2022.docx"],
        ["PR 决策事件（带 chosen_route_id）", "outputs/decisions/decision_events.parquet"],
        ["静态异构图", "outputs/static_graph/{nodes_*, edges_*}.parquet"],
        ["事件 token 流", "outputs/event_stream/event_tokens.parquet"],
        ["奖励数据", "outputs/rewards/decision_rewards.parquet"],
        ["奖励校准", "outputs/rewards/calibration.json"],
    ],
    col_widths_cm=[5.5, 10.0],
)

add_heading(doc, "B.4 关键源代码", 2)
add_table(doc,
    header=["模块", "路径"],
    rows=[
        ["路径配置", "src/railrl/config.py"],
        ["headcode / route_id 解析", "src/railrl/parsers.py"],
        ["数据 I/O 缓存", "src/railrl/data_io.py"],
        ["静态图构造", "src/railrl/data/static_graph.py"],
        ["事件流构造", "src/railrl/data/event_stream.py"],
        ["奖励校准", "src/railrl/data/reward_calibration.py"],
        ["奖励模型", "src/railrl/data/reward_model.py"],
        ["MDP 形式化（待写）", "src/railrl/mdp/"],
        ["编码器（待写）", "src/railrl/encoders/"],
        ["Q 网络 + 辅助头（待写）", "src/railrl/policies/"],
        ["训练算法（待写）", "src/railrl/algorithms/"],
        ["5 层 XAI（待写）", "src/railrl/xai/"],
    ],
    col_widths_cm=[5.5, 10.0],
)

doc.add_page_break()

# ============================================================
# Appendix C — Key numbers
# ============================================================
add_heading(doc, "附录 C   关键数字一览（一目了然）", 1)

add_heading(doc, "C.1 数据规模", 2)
add_table(doc,
    header=["指标", "数值"],
    rows=[
        ["数据时间范围", "2023-02-28 至 2024-04-25（14 个月）"],
        ["TD 原始事件", "11.91M 行（713 MB）"],
        ["TRUST 原始事件", "247k 行（49 MB）"],
        ["Panel_Request 决策总数", "546,418"],
        ["独立列车 ID", "2,185"],
        ["独立 signals 出现过", "92"],
        ["独立 routes 出现过", "250"],
    ],
    col_widths_cm=[7.5, 8.0],
)

add_heading(doc, "C.2 Derby 基础设施", 2)
add_table(doc,
    header=["类型", "数量"],
    rows=[
        ["Signals", "123"],
        ["Track sections（panel 内）", "249"],
        ["Routes（named）", "277"],
        ["Auxiliary connections", "156"],
        ["Edges（静态图 6 种总和）", "100 + 548 + 1701 + 279 + 290 + 1122"],
        ["TRTS（站台 ready 按钮）", "24（每 platform 4 个）"],
    ],
    col_widths_cm=[7.5, 8.0],
)

add_heading(doc, "C.3 决策分布", 2)
add_table(doc,
    header=["分类维度", "分布"],
    rows=[
        ["By prefix", "DC 195k / TD 145k / DW 121k / DY 58k / EC 27k"],
        ["By route class", "M 530k / S 14k / C 2.4k / SP 1"],
        ["By headcode class", "1=365k / 2=78k / 5=62k / 6=22k / 0=8k / 4=4.6k / 3=4.1k / 其他<1k"],
        ["Headcode parse rate", "99.49%"],
    ],
    col_widths_cm=[5.0, 10.5],
)

add_heading(doc, "C.4 Reward 校准参数（P2.4 经验值）", 2)
add_table(doc,
    header=["参数", "值", "百分位说明"],
    rows=[
        ["H_min（最小 headway）", "147.0 s", "P5 of 3.28M pair-wise headways"],
        ["d-gate 0.5 boundary", "6 hops", "P50 of 24,870 decisions"],
        ["d-gate 0.1 boundary", "16 hops", "P90"],
        ["Reward window", "4201.9 s", "P99 of 42,806 TIPLOC lags"],
        ["w_delay / w_throughput / w_headway / w_wait", "1.0 / 0.5 / 1.0 / 0.3", "默认权重"],
    ],
    col_widths_cm=[5.0, 3.5, 7.0],
)

add_heading(doc, "C.5 Reward 健康检查（2.64M 决策 / 82,429 episodes）", 2)
add_table(doc,
    header=["指标", "值"],
    rows=[
        ["r_total mean", "+0.255"],
        ["r_total std", "0.675"],
        ["per-episode return mean", "+2.25"],
        ["positive episode 比例", "88.4%"],
        ["Weight 排序稳健性（conservative vs default Spearman）", "0.908"],
        ["Movements proxy（Spearman r_total vs delay_reduction）", "+0.29"],
    ],
    col_widths_cm=[9.0, 6.5],
)

add_heading(doc, "C.6 实证审计关键比例", 2)
add_table(doc,
    header=["发现", "比例"],
    rows=[
        ["95.22% PR 触发时 conflict mask 内 TC 完全没被占用", "—"],
        ["2.19% PR 触发时被他车占用（大多 advance routing）", "—"],
        ["真实 PR 违反教科书 L1 first-TC-clear 规则", "3.01%"],
        ["非标准 train_id（343R 类）", "1.04%"],
        ["v1 per-signal majority binary baseline", "91%"],
        ["v1 B1 BC-MLP 准确率", "92.69%（仅比 trivial 高 1.4 pp）"],
    ],
    col_widths_cm=[12.0, 3.5],
)

doc.add_page_break()

# ============================================================
# Final notes
# ============================================================
add_heading(doc, "结语：如何使用本文档", 1)
add_para(doc, "当你开始一个新的对话讨论 RailRL v2 项目时，建议这样向 AI 助手介绍：")

add_callout(doc, "新对话开场建议",
"\"请先完整读 E:\\Claude\\RailRL_v2\\docs\\PROJECT_HANDOFF.docx，"
"特别是第 2 章（信号员的真实工作）、第 8 章（信息泄露契约）、第 11 章（v1 失败教训）。"
"然后再回答我接下来的问题。所有架构决策已锁定，不要建议改 task framing。"
"如果我让你写代码，先确认对应的 spec 是否已经写好——没写好的话先写 spec。\"",
color="E1EFFF")

add_para(doc, "如果你需要调整或扩展本文档，重新运行 docs/_build_handoff.py 即可。", italic=True)

# ============================================================
# Save
# ============================================================
output_path = '/sessions/epic-dazzling-feynman/mnt/Claude/RailRL_v2/docs/PROJECT_HANDOFF.docx'
doc.save(output_path)
print(f"✓ Saved to: {output_path}")
import os
print(f"  Size: {os.path.getsize(output_path) / 1024:.1f} KB")
print(f"  Paragraphs: {len(doc.paragraphs)}")
print(f"  Tables: {len(doc.tables)}")
")
")
print(f"  Tables: {len(doc.tables)}")
